"""
Exam Extraction Service for automatic question and answer extraction.
Processes uploaded questionnaire and answer key files to create exam questions.
"""
import logging
import re
from typing import List, Dict, Any, Optional, Tuple
from django.core.files.uploadedfile import UploadedFile
from services.file_service import FileService
from services.result import Result
from services.errors import ValidationError
from exams.models import Question, QuestionType

logger = logging.getLogger('services')


class ExamExtractionService:
    """
    Service for extracting questions and answers from uploaded exam files.
    Supports PDF, DOCX, and TXT formats.
    """
    
    def __init__(self):
        self.file_service = FileService()
    
    def extract_text_from_file(self, file: UploadedFile) -> Result[str, ValidationError]:
        """
        Extract text content from uploaded file.
        
        Args:
            file: Uploaded file (PDF, DOCX, or TXT)
            
        Returns:
            Result with extracted text or error
        """
        try:
            file_ext = file.name.lower().split('.')[-1]
            
            if file_ext == 'txt':
                # Read text file directly
                content = file.read().decode('utf-8', errors='ignore')
                return Result.success(content)
            
            elif file_ext == 'pdf':
                # Extract text from PDF using PyPDF2
                try:
                    import PyPDF2
                    import io
                    
                    # Reset file pointer
                    file.seek(0)
                    
                    # Create PDF reader
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
                    
                    # Extract text from all pages
                    text_content = []
                    for page in pdf_reader.pages:
                        text_content.append(page.extract_text())
                    
                    content = '\n'.join(text_content)
                    logger.info(f"Successfully extracted {len(content)} characters from PDF")
                    return Result.success(content)
                
                except ImportError:
                    logger.error("PyPDF2 library not installed")
                    return Result.failure(ValidationError(
                        message="PDF extraction requires PyPDF2 library. Please install: pip install PyPDF2",
                        details={'file_type': file_ext}
                    ))
                except Exception as e:
                    logger.error(f"Error extracting PDF: {str(e)}")
                    return Result.failure(ValidationError(
                        message=f"Failed to extract PDF content: {str(e)}",
                        details={'file_type': file_ext, 'error': str(e)}
                    ))
            
            elif file_ext in ['docx', 'doc']:
                # Extract text from DOCX using python-docx
                try:
                    from docx import Document
                    import io
                    
                    # Reset file pointer
                    file.seek(0)
                    
                    # Create Document object
                    doc = Document(io.BytesIO(file.read()))
                    
                    # Extract text from all paragraphs
                    text_content = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(paragraph.text)
                    
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    text_content.append(cell.text)
                    
                    content = '\n'.join(text_content)
                    logger.info(f"Successfully extracted {len(content)} characters from DOCX")
                    return Result.success(content)
                
                except ImportError:
                    logger.error("python-docx library not installed")
                    return Result.failure(ValidationError(
                        message="DOCX extraction requires python-docx library. Please install: pip install python-docx",
                        details={'file_type': file_ext}
                    ))
                except Exception as e:
                    logger.error(f"Error extracting DOCX: {str(e)}")
                    return Result.failure(ValidationError(
                        message=f"Failed to extract DOCX content: {str(e)}",
                        details={'file_type': file_ext, 'error': str(e)}
                    ))
            
            else:
                return Result.failure(ValidationError(
                    message=f"Unsupported file type: {file_ext}",
                    details={'file_type': file_ext}
                ))
        
        except Exception as e:
            logger.error(f"Error extracting text from file: {str(e)}")
            return Result.failure(ValidationError(
                message="Failed to extract text from file",
                details={'error': str(e)}
            ))
    
    def parse_questions_from_text(self, text: str) -> List[Dict[str, Any]]:
        """
        Parse questions from extracted text.
        Detects question patterns and extracts question data.
        
        Supported formats:
        - Numbered questions: "1. Question text?" or "1) Question text?"
        - Unnumbered questions (auto-numbered by parser)
        - Section headers: "Multiple Choice (2 points each)"
        - MCQ with inline options: "A. Option1 B. Option2 C. Option3"
        
        Args:
            text: Extracted text content
            
        Returns:
            List of question dictionaries
        """
        questions = []
        
        # Split text into lines
        lines = text.split('\n')
        
        current_question = None
        current_options = []
        question_number = 0
        current_section = None
        auto_number = 1
        
        # Section headers to skip
        skip_patterns = [
            r'^Grade \d+',
            r'^Total Items:',
            r'^Types:',
            r'^Answer Key',
            r'^\(Answers may vary\)',
            r'^Identification \(continued\)',
            r'^Multiple Choice \(continued\)',
        ]
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Skip header lines
            if any(re.match(pattern, line, re.IGNORECASE) for pattern in skip_patterns):
                continue
            
            # Check for section headers
            section_match = re.match(r'^(Identification|Multiple Choice|Enumeration|Essay|True.*False)\s*(\(.*\))?', line, re.IGNORECASE)
            if section_match:
                current_section = section_match.group(1)
                logger.info(f"Found section: {current_section}")
                continue
            
            # Check if line starts with a number (explicit question number)
            question_match = re.match(r'^(\d+)\s*[\.\):\-]\s*(.+)$', line)
            
            if question_match:
                # Save previous question if exists
                if current_question:
                    questions.append(self._finalize_question(
                        question_number if question_number > 0 else auto_number,
                        current_question,
                        current_options
                    ))
                    auto_number += 1
                
                # Start new question
                question_number = int(question_match.group(1))
                current_question = question_match.group(2).strip()
                current_options = []
                
                # Check if options are on the same line
                options_in_line = self._extract_inline_options(current_question)
                if options_in_line:
                    current_question = options_in_line['question']
                    current_options = options_in_line['options']
            
            # Check if line is an MCQ option (A., B., C., etc.)
            elif re.match(r'^([A-D])\s*[\.\)]\s+', line):
                option_match = re.match(r'^([A-D])\s*[\.\)]\s+(.+)$', line)
                if option_match:
                    # If we have options, this might be a new question
                    if not current_question:
                        # This is the start of an MCQ question
                        # The question text might be on the previous line or missing
                        pass
                    
                    option_key = option_match.group(1).upper()
                    option_value = option_match.group(2).strip()
                    current_options.append({
                        'key': option_key,
                        'value': option_value
                    })
            
            # Otherwise, treat as a question (for unnumbered formats)
            else:
                # Check if this looks like a question
                # Skip very short lines or lines that look like section markers
                if len(line) > 10 and not re.match(r'^[A-Z\s\(\)]{5,}$', line):
                    # If we have a current question with options, save it first
                    if current_question and current_options:
                        questions.append(self._finalize_question(
                            question_number if question_number > 0 else auto_number,
                            current_question,
                            current_options
                        ))
                        auto_number += 1
                        current_question = None
                        current_options = []
                        question_number = 0
                    
                    # Start new question or append to current
                    if not current_question:
                        current_question = line
                    else:
                        # Might be continuation of previous question
                        # Only append if it doesn't look like an option
                        if not re.match(r'^[A-D][\.\)]', line):
                            current_question += ' ' + line
        
        # Add last question
        if current_question:
            questions.append(self._finalize_question(
                question_number if question_number > 0 else auto_number,
                current_question,
                current_options
            ))
        
        logger.info(f"Extracted {len(questions)} questions from text")
        return questions
    
    def _extract_inline_options(self, text: str) -> Optional[Dict[str, Any]]:
        """
        Extract options that are inline with the question.
        Supports formats:
        - "What is 2+2? A) 3 B) 4 C) 5"
        - "What is 2+2? A. 3 B. 4 C. 5"
        """
        # Try pattern with parentheses first: A) B) C) D)
        pattern1 = r'([A-D])\)\s*([^A-D\)]+?)(?=\s+[A-D]\)|$)'
        matches = re.findall(pattern1, text)
        
        if matches and len(matches) >= 2:
            # Extract question text (before first option)
            question_text = re.split(r'\s+[A-D][\.\)]', text)[0].strip()
            
            options = []
            for key, value in matches:
                options.append({
                    'key': key.upper(),
                    'value': value.strip()
                })
            
            return {
                'question': question_text,
                'options': options
            }
        
        # Try pattern with periods: A. B. C. D.
        pattern2 = r'([A-D])\.\s+([^A-D\.]+?)(?=\s+[A-D]\.|$)'
        matches = re.findall(pattern2, text)
        
        if matches and len(matches) >= 2:
            # Extract question text (before first option)
            question_text = re.split(r'\s+[A-D]\.', text)[0].strip()
            
            options = []
            for key, value in matches:
                # Clean up the value
                value = value.strip()
                # Remove trailing period if present
                if value.endswith('.'):
                    value = value[:-1].strip()
                options.append({
                    'key': key.upper(),
                    'value': value
                })
            
            return {
                'question': question_text,
                'options': options
            }
        
        return None
    
    def _finalize_question(
        self,
        number: int,
        question_text: str,
        options: List[Dict[str, str]]
    ) -> Dict[str, Any]:
        """
        Finalize question data structure.
        """
        # Determine question type
        if options and len(options) >= 2:
            question_type = QuestionType.MCQ
        elif 'true or false' in question_text.lower() or 't/f' in question_text.lower():
            question_type = QuestionType.TRUE_FALSE
        elif 'enumerate' in question_text.lower() or 'list' in question_text.lower():
            question_type = QuestionType.ENUMERATION
        elif 'identify' in question_text.lower() or 'what is' in question_text.lower():
            question_type = QuestionType.IDENTIFICATION
        else:
            question_type = QuestionType.ESSAY
        
        return {
            'number': number,
            'question_type': question_type,
            'question_text': question_text,
            'options': options if options else None,
            'points': 1.0,
            'order_index': number
        }
    
    def debug_extracted_text(self, text: str, max_lines: int = 50) -> str:
        """
        Return first N lines of extracted text for debugging.
        
        Args:
            text: Extracted text
            max_lines: Maximum number of lines to return
            
        Returns:
            First N lines of text
        """
        lines = text.split('\n')[:max_lines]
        return '\n'.join(f"{i+1}: {line}" for i, line in enumerate(lines))
    
    def parse_answers_from_text(self, text: str) -> Dict[int, Any]:
        """
        Parse answer key from extracted text.
        
        Supported formats:
        - "1. A" or "1) A" for MCQ
        - "1. Answer text" for identification
        - "1. True" or "1. False" for T/F
        
        Args:
            text: Extracted answer key text
            
        Returns:
            Dictionary mapping question number to answer
        """
        answers = {}
        
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Match pattern: "1. A" or "1) Answer"
            match = re.match(r'^(\d+)[\.\)]\s*(.+)$', line)
            
            if match:
                question_num = int(match.group(1))
                answer_text = match.group(2).strip()
                
                # Determine answer type
                if len(answer_text) == 1 and answer_text.isalpha():
                    # Single letter - likely MCQ answer
                    answers[question_num] = answer_text.upper()
                elif answer_text.lower() in ['true', 'false', 't', 'f']:
                    # True/False answer
                    answers[question_num] = answer_text.lower() in ['true', 't']
                else:
                    # Text answer - could be identification or enumeration
                    # Check if it contains commas (enumeration)
                    if ',' in answer_text:
                        items = [item.strip() for item in answer_text.split(',')]
                        answers[question_num] = {
                            'items': items,
                            'min_required': len(items)
                        }
                    else:
                        # Single text answer
                        answers[question_num] = [answer_text]
        
        logger.info(f"Extracted {len(answers)} answers from answer key")
        return answers
    
    def merge_questions_and_answers(
        self,
        questions: List[Dict[str, Any]],
        answers: Dict[int, Any]
    ) -> List[Dict[str, Any]]:
        """
        Merge extracted questions with their answers.
        
        Args:
            questions: List of extracted questions
            answers: Dictionary of answers keyed by question number
            
        Returns:
            List of complete question data with answers
        """
        merged = []
        
        for question in questions:
            question_num = question.get('number', 0)
            
            if question_num in answers:
                question['correct_answer'] = answers[question_num]
            else:
                # No answer found - set to None for manual entry
                question['correct_answer'] = None
            
            merged.append(question)
        
        logger.info(f"Merged {len(merged)} questions with answers")
        return merged
    
    def create_questions_from_extracted_data(
        self,
        exam,
        extracted_questions: List[Dict[str, Any]]
    ) -> List[Question]:
        """
        Create Question objects from extracted data.
        
        Args:
            exam: Exam instance to attach questions to
            extracted_questions: List of extracted question data
            
        Returns:
            List of created Question instances
        """
        created_questions = []
        
        for q_data in extracted_questions:
            try:
                question = Question.objects.create(
                    exam=exam,
                    question_type=q_data['question_type'],
                    question_text=q_data['question_text'],
                    options=q_data.get('options'),
                    correct_answer=q_data.get('correct_answer'),
                    points=q_data.get('points', 1.0),
                    order_index=q_data.get('order_index', 0)
                )
                created_questions.append(question)
                logger.info(f"Created question {question.id} for exam {exam.id}")
            
            except Exception as e:
                logger.error(f"Failed to create question: {str(e)}")
                continue
        
        return created_questions
