from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.http import JsonResponse
from django.views.decorators.http import require_http_methods
from django.views.decorators.csrf import ensure_csrf_cookie
from django.utils import timezone
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db.models import Count, Avg
from exams.models import Exam, Question, QuestionType, AIGenerationTask
from exams.forms import ExamForm, QuestionForm
from users.models import Class
from users.decorators import teacher_required
from services.auth_service import AuthenticationService
from services.exam_service import ExamService
from services.question_service import QuestionService
from services.exam_activation_service import ExamActivationService
from services.view_helpers import build_breadcrumbs
from django.urls import reverse
import logging
import re
import json
from itertools import groupby

logger = logging.getLogger(__name__)


# Initialize services
exam_service = ExamService()
question_service = QuestionService()
activation_service = ExamActivationService()
auth_service = AuthenticationService()


@teacher_required
def exam_list_view(request):
    """
    Display all exams for the logged-in teacher with pagination.
    Requirements: 2.5, 8.1
    """
    # Get teacher profile
    teacher = auth_service.get_current_teacher(request)
    
    # Get all exams created by this teacher
    exams = exam_service.get_exams_by_teacher(teacher.pk)
    
    # Implement pagination (20 items per page)
    paginator = Paginator(exams, 20)
    page_number = request.GET.get('page', 1)
    
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        # If page is not an integer, deliver first page
        page_obj = paginator.page(1)
    except EmptyPage:
        # If page is out of range, deliver last page
        page_obj = paginator.page(paginator.num_pages)

    page_exams = list(page_obj.object_list)
    page_exams.sort(key=lambda exam: (
        exam.quarter.order if exam.quarter else 9999,
        exam.quarter.name.lower() if exam.quarter else 'no quarter',
        -exam.created_at.timestamp(),
    ))

    grouped_exams = []
    for key, group in groupby(
        page_exams,
        key=lambda exam: (
            exam.quarter.id if exam.quarter else None,
            exam.quarter.name if exam.quarter else 'No Quarter',
            exam.quarter.order if exam.quarter else 9999,
        )
    ):
        group_items = list(group)
        grouped_exams.append({
            'quarter_id': key[0],
            'label': key[1],
            'exams': group_items,
        })
    
    # Build breadcrumbs
    breadcrumbs = build_breadcrumbs(
        ('Dashboard', reverse('teacher_dashboard')),
        'My Exams'
    )
    
    context = {
        'exams': page_obj,
        'grouped_exams': grouped_exams,
        'page_obj': page_obj,
        'teacher': teacher,
        'page_breadcrumbs': breadcrumbs
    }
    return render(request, 'exams/exam_list.html', context)


@teacher_required
def exam_create_view(request):
    """
    Create a new exam with form validation and support for two creation methods:
    1. File upload - Extract questions from uploaded documents
    2. Manual entry - Create questions manually in the editor
    
    Requirements: 1.5, 3.1
    """
    teacher = auth_service.get_current_teacher(request)
    
    if request.method == 'POST':
        form = ExamForm(request.POST, request.FILES, teacher=teacher)
        is_ajax = request.headers.get('X-Requested-With') == 'XMLHttpRequest'

        if not form.is_valid():
            if is_ajax:
                errors = [e for field_errors in form.errors.values() for e in field_errors]
                return JsonResponse({'success': False, 'error': errors[0] if errors else 'Form validation failed.'})

            # Display field-specific error messages (Requirement 1.4)
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, error)
            
            # Get teacher's classes for the template
            classes = Class.objects.filter(teachers=teacher).order_by('grade_level', 'strand', 'section')
            
            # Build breadcrumbs
            breadcrumbs = build_breadcrumbs(
                ('Dashboard', reverse('teacher_dashboard')),
                ('My Exams', reverse('exam_list')),
                'Create Exam'
            )
            
            return render(request, 'exams/exam_form.html', {
                'form': form,
                'classes': classes,
                'page_breadcrumbs': breadcrumbs
            })
        
        # Get generation method to route appropriately
        generation_method = form.cleaned_data.get('generation_method', 'manual')
        
        # Create exam first
        exam_data = {
            'title': form.cleaned_data['title'],
            'subject': form.cleaned_data.get('subject', ''),
            'quarter': form.cleaned_data.get('quarter'),
            'description': form.cleaned_data.get('description', ''),
            'duration_minutes': form.cleaned_data['duration_minutes'],
            'created_by': teacher
        }
        
        exam = exam_service.create_exam(exam_data)
        if not exam:
            if is_ajax:
                return JsonResponse({'success': False, 'error': 'Failed to create exam.'})
            messages.error(request, 'Failed to create exam')
            
            # Get teacher's classes for the template
            classes = Class.objects.filter(teachers=teacher).order_by('grade_level', 'strand', 'section')
            
            # Build breadcrumbs
            breadcrumbs = build_breadcrumbs(
                ('Dashboard', reverse('teacher_dashboard')),
                ('My Exams', reverse('exam_list')),
                'Create Exam'
            )
            
            return render(request, 'exams/exam_form.html', {
                'form': form,
                'classes': classes,
                'page_breadcrumbs': breadcrumbs
            })
        
        # Create ExamClassAssignment records (Requirement 3.1)
        assigned_classes = form.cleaned_data.get('assigned_classes', [])
        if assigned_classes:
            from exams.models import ExamClassAssignment
            for class_obj in assigned_classes:
                ExamClassAssignment.objects.create(
                    exam=exam,
                    class_assigned=class_obj
                )
            messages.info(
                request,
                f'Exam assigned to {len(assigned_classes)} class(es)'
            )
        
        # Route to appropriate service based on generation method
        if generation_method == 'ai_generate':
            # AI generation - start background task
            import threading
            is_ajax = request.headers.get('X-Requested-With') == 'XMLHttpRequest'

            try:
                topic = request.POST.get('ai_topic', '').strip()
                difficulty = request.POST.get('ai_difficulty', 'medium')
                subject = exam_data.get('subject', '')
                grade_level = request.POST.get('ai_grade_level', 'grade_11_12')

                # Build per-type counts (safe int parsing)
                type_counts = {}
                for qt in ['MCQ', 'TRUE_FALSE', 'IDENTIFICATION', 'ENUMERATION', 'ESSAY']:
                    raw = request.POST.get(f'ai_count_{qt}', '0')
                    try:
                        count = int(raw) if raw else 0
                    except (ValueError, TypeError):
                        count = 0
                    if count > 0:
                        type_counts[qt] = count

                if not topic:
                    if is_ajax:
                        return JsonResponse({'success': False, 'error': 'No topic provided for AI generation.'})
                    messages.warning(request, 'No topic provided for AI generation.')
                    return redirect('exam_edit', exam_id=exam.id)

                if not type_counts:
                    if is_ajax:
                        return JsonResponse({'success': False, 'error': 'No question counts specified.'})
                    messages.warning(request, 'No question counts specified.')
                    return redirect('exam_edit', exam_id=exam.id)

                # Create task record
                task = AIGenerationTask.objects.create(
                    exam=exam,
                    topic=topic,
                    subject=subject,
                    difficulty=difficulty,
                    type_counts=type_counts,
                    total_requested=sum(type_counts.values()),
                    status='pending',
                )

                # Start background thread
                thread = threading.Thread(
                    target=_run_ai_generation,
                    args=(task.id, grade_level),
                    daemon=True,
                )
                thread.start()

                if is_ajax:
                    return JsonResponse({
                        'success': True,
                        'task_id': task.id,
                        'exam_id': exam.id,
                        'status': 'pending',
                    })

                messages.info(request, 'AI generation started. Questions will appear shortly.')
                return redirect('exam_edit', exam_id=exam.id)

            except Exception as e:
                import traceback
                import logging
                logging.getLogger(__name__).error(f'AI generation error: {traceback.format_exc()}')
                if is_ajax:
                    return JsonResponse({'success': False, 'error': f'Server error: {str(e)}'})
                messages.error(request, f'AI generation error: {e}')
                return redirect('exam_edit', exam_id=exam.id)

        else:  # manual
            # Manual entry - redirect to exam editor
            messages.success(
                request,
                f'Exam "{exam.title}" created successfully. Add questions manually.'
            )
            return redirect('exam_edit', exam_id=exam.id)
    
    form = ExamForm(teacher=teacher)
    
    # Get teacher's classes for the template
    classes = Class.objects.filter(teachers=teacher).order_by('grade_level', 'strand', 'section')
    
    # Build breadcrumbs
    breadcrumbs = build_breadcrumbs(
        ('Dashboard', reverse('teacher_dashboard')),
        ('My Exams', reverse('exam_list')),
        'Create Exam'
    )
    
    return render(request, 'exams/exam_form.html', {
        'form': form,
        'classes': classes,
        'page_breadcrumbs': breadcrumbs
    })


def _run_ai_generation(task_id, grade_level='grade_11_12'):
    """Background worker that generates AI questions in batches."""
    import django
    django.db.connections.close_all()

    from services.ai_generation_service import generate_exam_questions

    try:
        task = AIGenerationTask.objects.get(pk=task_id)
        task.status = 'processing'
        task.save(update_fields=['status'])

        questions = generate_exam_questions(
            task.topic, task.subject,
            type_counts=task.type_counts,
            difficulty=task.difficulty,
            grade_level=grade_level,
        )

        questions_by_type = {}
        for i, q in enumerate(questions):
            correct_answer = q.get('correct_answer')
            if correct_answer is None:
                correct_answer = ''
            Question.objects.create(
                exam=task.exam,
                question_type=q['question_type'],
                question_text=q['question_text'],
                options=q.get('options', []),
                correct_answer=correct_answer,
                points=q.get('points', 1.0),
                order_index=i + 1,
            )
            qt = q['question_type']
            questions_by_type[qt] = questions_by_type.get(qt, 0) + 1

        task.status = 'completed'
        task.total_generated = len(questions)
        task.questions_by_type = questions_by_type
        task.completed_at = timezone.now()
        task.save()

    except Exception as e:
        try:
            task = AIGenerationTask.objects.get(pk=task_id)
            task.status = 'failed'
            task.error_message = str(e)
            task.completed_at = timezone.now()
            task.save()
        except Exception:
            pass


@teacher_required
@require_http_methods(["GET"])
def ai_task_status_view(request, task_id):
    """Polling endpoint to check AI generation task status."""
    task = get_object_or_404(AIGenerationTask, pk=task_id)
    teacher = auth_service.get_current_teacher(request)
    if task.exam.created_by.pk != teacher.pk:
        return JsonResponse({'error': 'Permission denied'}, status=403)

    return JsonResponse({
        'task_id': task.id,
        'exam_id': task.exam.id,
        'status': task.status,
        'total_requested': task.total_requested,
        'total_generated': task.total_generated,
        'questions_by_type': task.questions_by_type,
        'error': task.error_message,
    })


@teacher_required
def exam_create_test_view(request):
    """
    Simple test view for exam creation form
    """
    teacher = auth_service.get_current_teacher(request)
    
    if request.method == 'POST':
        # Just redirect back for testing
        messages.success(request, 'Test form submitted successfully!')
        return redirect('exam_create_test')
    
    # Get teacher's classes for the template
    from users.models import Class
    classes = Class.objects.filter(teachers=teacher).order_by('grade_level', 'strand', 'section')
    
    return render(request, 'exams/exam_form_simple.html', {
        'classes': classes
    })



@teacher_required
def exam_edit_view(request, exam_id):
    """
    Edit an existing exam and manage its questions.
    
    This view handles editing for exams created through two methods:
    - File upload
    - Manual entry
    
    All questions can be edited regardless of creation method (Requirement 6.5).
    
    Requirements: 1.2, 6.5, 3.1, 3.5
    """
    exam = get_object_or_404(Exam, pk=exam_id)
    teacher = auth_service.get_current_teacher(request)
    
    # Check if user owns this exam
    if exam.created_by.pk != teacher.pk:
        messages.error(request, 'You do not have permission to edit this exam')
        return redirect('exam_list')
    
    if request.method == 'POST':
        form = ExamForm(request.POST, instance=exam, teacher=teacher)
        
        if not form.is_valid():
            # Display field-specific error messages
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, error)
        else:
            # Update exam with validated data
            exam_data = {
                'title': form.cleaned_data['title'],
                'subject': form.cleaned_data.get('subject', ''),
                'quarter': form.cleaned_data.get('quarter'),
                'description': form.cleaned_data.get('description', ''),
                'duration_minutes': form.cleaned_data['duration_minutes']
            }
            updated_exam = exam_service.update_exam(exam_id, exam_data)
            if updated_exam:
                # Update ExamClassAssignment records (Requirements 3.1, 3.5)
                from exams.models import ExamClassAssignment
                assigned_classes = form.cleaned_data.get('assigned_classes', [])
                
                # Get current assignments
                current_assignments = set(
                    ExamClassAssignment.objects.filter(exam=exam).values_list('class_assigned_id', flat=True)
                )
                new_assignments = set(c.id for c in assigned_classes)
                
                # Remove assignments that are no longer selected
                to_remove = current_assignments - new_assignments
                if to_remove:
                    ExamClassAssignment.objects.filter(
                        exam=exam,
                        class_assigned_id__in=to_remove
                    ).delete()
                
                # Add new assignments
                to_add = new_assignments - current_assignments
                for class_id in to_add:
                    ExamClassAssignment.objects.create(
                        exam=exam,
                        class_assigned_id=class_id
                    )
                
                messages.success(request, 'Exam updated successfully')
                exam = updated_exam
            else:
                messages.error(request, 'Failed to update exam')
    else:
        # Pre-populate assigned_classes with current assignments
        from exams.models import ExamClassAssignment
        current_class_ids = ExamClassAssignment.objects.filter(
            exam=exam
        ).values_list('class_assigned_id', flat=True)
        
        form = ExamForm(instance=exam, teacher=teacher)
        form.initial['assigned_classes'] = list(current_class_ids)
    
    # Get questions for this exam
    questions = question_service.get_questions_by_exam(exam_id)
    
    # Serialize question data for JavaScript
    questions_data = []
    for q in questions:
        # Get options - check if it's a field or property
        options_value = None
        if hasattr(q, 'options') and q.options:
            options_value = q.options
        
        # Get correct answer
        correct_answer_value = None
        if hasattr(q, 'correct_answer') and q.correct_answer:
            correct_answer_value = q.correct_answer
        
        q_data = {
            'id': q.id,
            'type': q.question_type,
            'text': q.question_text,
            'points': float(q.points),
            'correct_answer': correct_answer_value,
            'options': options_value,
        }
        questions_data.append(q_data)
    
    logger.info(f"Serialized {len(questions_data)} questions for JavaScript")
    if questions_data:
        logger.info(f"Sample question data: {questions_data[0]}")
    
    # Get current class assignments for display (Requirement 3.1)
    from exams.models import ExamClassAssignment
    from users.models import Class
    assigned_classes = Class.objects.filter(
        exam_assignments__exam=exam
    ).order_by('grade_level', 'strand', 'section')
    
    # Get all teacher's classes for the form
    classes = Class.objects.filter(teachers=teacher).order_by('grade_level', 'strand', 'section')
    
    # Build breadcrumbs
    breadcrumbs = build_breadcrumbs(
        ('Dashboard', reverse('teacher_dashboard')),
        ('My Exams', reverse('exam_list')),
        f'Edit: {exam.title}'
    )
    
    context = {
        'exam': exam,
        'form': form,
        'questions': questions,
        'questions_data': questions_data,
        'question_types': QuestionType.choices,
        'assigned_classes': assigned_classes,
        'classes': classes,
        'page_breadcrumbs': breadcrumbs
    }
    return render(request, 'exams/exam_edit.html', context)


@teacher_required
def exam_detail_view(request, exam_id):
    """
    Display a read-only overview of an exam with quick actions,
    question list, assigned classes, and taker summary.
    """
    exam = get_object_or_404(Exam, pk=exam_id)
    teacher = auth_service.get_current_teacher(request)

    if exam.created_by.pk != teacher.pk:
        messages.error(request, 'You do not have permission to view this exam')
        return redirect('exam_list')

    from attempts.models import Attempt, AttemptStatus
    from exams.models import ExamStudentAssignment
    from users.models import Student

    questions = list(question_service.get_questions_by_exam(exam_id))
    question_count = len(questions)
    total_points = sum(float(question.points) for question in questions)

    assigned_classes = Class.objects.filter(
        exam_assignments__exam=exam
    ).annotate(
        student_count=Count('students', distinct=True)
    ).order_by('grade_level', 'strand', 'section')

    assigned_class_count = assigned_classes.count()
    assigned_class_ids = list(assigned_classes.values_list('id', flat=True))
    assigned_students_count = Student.objects.filter(
        class_assigned_id__in=assigned_class_ids
    ).count() if assigned_class_ids else 0

    student_assignments = list(
        ExamStudentAssignment.objects.filter(exam=exam).select_related(
            'student', 'student__class_assigned'
        ).order_by('student__last_name', 'student__first_name')
    )
    access_mode = 'Assigned classes'
    accessible_students_count = assigned_students_count
    selected_students = []
    if student_assignments:
        access_mode = 'Selected students'
        accessible_students_count = len(student_assignments)
        selected_students = [assignment.student for assignment in student_assignments]

    attempts = Attempt.objects.filter(
        exam=exam
    ).select_related(
        'student', 'student__class_assigned'
    ).order_by('-submitted_at', '-started_at')

    total_attempts = attempts.count()
    in_progress_attempts = attempts.filter(status=AttemptStatus.IN_PROGRESS).count()
    submitted_attempts = attempts.filter(status=AttemptStatus.SUBMITTED).count()
    graded_attempts = attempts.filter(status=AttemptStatus.GRADED).count()
    unique_takers = attempts.filter(
        status__in=[AttemptStatus.SUBMITTED, AttemptStatus.GRADED]
    ).values('student_id').distinct().count()
    flagged_attempts = attempts.filter(is_flagged=True).count()

    average_score = attempts.filter(
        status__in=[AttemptStatus.SUBMITTED, AttemptStatus.GRADED]
    ).aggregate(avg_score=Avg('total_score'))['avg_score']
    average_score = float(average_score) if average_score is not None else None

    recent_attempts_data = []
    for attempt in attempts[:8]:
        percentage = 0
        if total_points > 0:
            percentage = (float(attempt.total_score) / total_points) * 100
        submitted_moment = attempt.submitted_at or attempt.started_at
        recent_attempts_data.append({
            'attempt': attempt,
            'student': attempt.student,
            'class_name': str(attempt.student.class_assigned) if attempt.student.class_assigned else 'No Class',
            'percentage': round(percentage, 2),
            'submitted_display': submitted_moment.strftime('%b %d, %Y %H:%M'),
        })

    breadcrumbs = build_breadcrumbs(
        ('Dashboard', reverse('teacher_dashboard')),
        ('My Exams', reverse('exam_list')),
        f'Exam Details: {exam.title}'
    )

    context = {
        'exam': exam,
        'questions': questions,
        'question_count': question_count,
        'total_points': total_points,
        'assigned_classes': assigned_classes,
        'assigned_class_count': assigned_class_count,
        'assigned_students_count': assigned_students_count,
        'selected_students': selected_students,
        'access_mode': access_mode,
        'accessible_students_count': accessible_students_count,
        'total_attempts': total_attempts,
        'in_progress_attempts': in_progress_attempts,
        'submitted_attempts': submitted_attempts,
        'graded_attempts': graded_attempts,
        'unique_takers': unique_takers,
        'flagged_attempts': flagged_attempts,
        'average_score': average_score,
        'recent_attempts_data': recent_attempts_data,
        'page_breadcrumbs': breadcrumbs,
    }

    return render(request, 'exams/exam_detail.html', context)


@teacher_required
@require_http_methods(["POST"])
def exam_activate_view(request, exam_id):
    """
    Toggle exam activation status.
    When reopening, allows selection of which students can access the exam.
    """
    exam = get_object_or_404(Exam, pk=exam_id)
    
    # Check if user owns this exam
    teacher = auth_service.get_current_teacher(request)
    if exam.created_by.pk != teacher.pk:
        messages.error(request, 'You do not have permission to modify this exam')
        return redirect('exam_list')
    
    # Check if we're reactivating (exam is currently inactive)
    was_inactive = not exam.is_active
    
    # Toggle activation
    updated_exam = activation_service.toggle_activation(exam_id)
    
    if updated_exam:
        if updated_exam.is_active and was_inactive:
            # Exam is being reopened - handle student selection
            from repositories.exam_student_assignment_repository import ExamStudentAssignmentRepository
            from exams.models import ExamClassAssignment
            from users.models import Student
            
            assignment_repo = ExamStudentAssignmentRepository()
            
            # Clear existing assignments (reset access)
            assignment_repo.clear_assignments_for_exam(exam_id)
            
            # Get student IDs from POST
            student_ids = request.POST.getlist('student_ids')
            select_all = request.POST.get('select_all') == '1'
            
            if select_all:
                # Get all students from assigned classes
                assigned_classes = ExamClassAssignment.objects.filter(
                    exam=exam
                ).values_list('class_assigned_id', flat=True)
                
                all_student_ids = list(
                    Student.objects.filter(
                        class_assigned_id__in=assigned_classes
                    ).values_list('id', flat=True)
                )
                
                if all_student_ids:
                    assignment_repo.assign_students_to_exam(exam_id, all_student_ids)
                    messages.success(
                        request,
                        f'Exam "{updated_exam.title}" reopened successfully. All students from assigned classes can now access it.'
                    )
                else:
                    messages.warning(
                        request,
                        f'Exam "{updated_exam.title}" reopened, but no students found in assigned classes.'
                    )
            elif student_ids:
                # Assign selected students
                try:
                    student_ids_int = [int(sid) for sid in student_ids]
                    assignment_repo.assign_students_to_exam(exam_id, student_ids_int)
                    messages.success(
                        request,
                        f'Exam "{updated_exam.title}" reopened successfully. {len(student_ids_int)} selected student(s) can now access it.'
                    )
                except (ValueError, TypeError):
                    messages.error(request, 'Invalid student selection. Please try again.')
            else:
                # No students selected - exam reopened but no one can access
                messages.warning(
                    request,
                    f'Exam "{updated_exam.title}" reopened, but no students were selected. The exam is active but not accessible to any students.'
                )
        elif updated_exam.is_active:
            # Exam was already active, just activated again (shouldn't happen normally)
            messages.success(
                request, 
                f'Exam "{updated_exam.title}" is already active.'
            )
        else:
            # Exam is being closed
            messages.success(
                request, 
                f'Exam "{updated_exam.title}" closed successfully. It is now hidden from students.'
            )
    else:
        messages.error(request, 'Failed to update exam status')
    
    return redirect('exam_list')


@teacher_required
@require_http_methods(["POST"])
def exam_delete_view(request, exam_id):
    """
    Delete an exam with password confirmation for enhanced security.
    
    This view requires the user to enter their password to confirm deletion
    since exams contain important data (questions, student attempts, grades).
    
    Requirements:
    - User must own the exam
    - User must provide correct password
    - Deletes exam and all related data (questions, attempts, grades)
    """
    exam = get_object_or_404(Exam, pk=exam_id)
    teacher = auth_service.get_current_teacher(request)
    
    # Check if user owns this exam
    if exam.created_by.pk != teacher.pk:
        messages.error(request, 'You do not have permission to delete this exam')
        return redirect('exam_list')
    
    # Get password from request
    password = request.POST.get('password', '')
    
    if not password:
        messages.error(request, 'Password is required to delete an exam')
        return redirect('exam_list')
    
    # Verify password
    from django.contrib.auth import authenticate
    user = authenticate(username=teacher.user.username, password=password)
    
    if user is None:
        messages.error(
            request,
            'Incorrect password. Exam deletion cancelled for security.'
        )
        return redirect('exam_list')
    
    # Password verified - proceed with deletion
    exam_title = exam.title
    
    try:
        # Get counts for confirmation message
        question_count = exam.questions.count()
        from attempts.models import Attempt
        attempt_count = Attempt.objects.filter(exam=exam).count()
        
        # Delete the exam (cascade will delete related questions, attempts, etc.)
        exam.delete()
        
        logger.info(
            f"Exam deleted: '{exam_title}' (id={exam_id}) by teacher {teacher.user.username}. "
            f"Deleted {question_count} questions and {attempt_count} attempts."
        )
        
        messages.success(
            request,
            f'Exam "{exam_title}" has been permanently deleted. '
            f'Removed {question_count} questions and {attempt_count} student attempts.'
        )
        
    except Exception as e:
        logger.error(f"Error deleting exam {exam_id}: {str(e)}", exc_info=True)
        messages.error(
            request,
            'An error occurred while deleting the exam. Please try again or contact support.'
        )
    
    return redirect('exam_list')


@teacher_required
@require_http_methods(["POST"])
def question_create_view(request, exam_id):
    """
    Create a new question for an exam.
    """
    exam = get_object_or_404(Exam, pk=exam_id)
    
    # Check if user owns this exam
    teacher = auth_service.get_current_teacher(request)
    if exam.created_by.pk != teacher.pk:
        return JsonResponse({'error': 'Permission denied'}, status=403)
    
    # Get form data
    question_type = request.POST.get('question_type')
    question_text = request.POST.get('question_text', '').strip()
    points = request.POST.get('points', '1.0')
    
    # Validation
    if not question_text:
        return JsonResponse({'error': 'Question text is required'}, status=400)
    
    try:
        points = float(points)
        if points <= 0:
            return JsonResponse({'error': 'Points must be positive'}, status=400)
    except ValueError:
        return JsonResponse({'error': 'Points must be a valid number'}, status=400)
    
    # Prepare question data based on type
    question_data = {
        'exam': exam,
        'question_type': question_type,
        'question_text': question_text,
        'points': points
    }
    
    # Handle type-specific data
    if question_type == QuestionType.MCQ:
        # Parse MCQ options
        options = []
        correct_answer = request.POST.get('correct_answer', '')
        option_keys = request.POST.getlist('option_keys[]')
        option_values = request.POST.getlist('option_values[]')
        
        for key, value in zip(option_keys, option_values):
            if key and value:
                options.append({'key': key, 'value': value})
        
        if len(options) < 2:
            return JsonResponse({'error': 'MCQ must have at least 2 options'}, status=400)
        if not correct_answer:
            return JsonResponse({'error': 'Correct answer is required'}, status=400)
        
        question_data['options'] = options
        question_data['correct_answer'] = correct_answer
    
    elif question_type == QuestionType.IDENTIFICATION:
        correct_answer = request.POST.get('correct_answer', '').strip()
        if not correct_answer:
            return JsonResponse({'error': 'Correct answer is required'}, status=400)
        # Store as list to support multiple acceptable answers
        question_data['correct_answer'] = [correct_answer]
    
    elif question_type == QuestionType.ENUMERATION:
        correct_answers = request.POST.get('correct_answers', '').strip()
        if not correct_answers:
            return JsonResponse({'error': 'Correct answers are required'}, status=400)
        # Split by comma or newline
        answers = [ans.strip() for ans in correct_answers.replace('\n', ',').split(',') if ans.strip()]
        min_required = request.POST.get('min_required', len(answers))
        try:
            min_required = int(min_required)
        except ValueError:
            min_required = len(answers)
        
        question_data['correct_answer'] = {
            'items': answers,
            'min_required': min_required
        }
    
    elif question_type == QuestionType.TRUE_FALSE:
        correct_answer = request.POST.get('correct_answer', '')
        if correct_answer not in ['true', 'false']:
            return JsonResponse({'error': 'Correct answer must be true or false'}, status=400)
        question_data['correct_answer'] = correct_answer == 'true'
    
    elif question_type == QuestionType.ESSAY:
        # Essay questions don't have a correct answer
        question_data['correct_answer'] = None
    
    # Create question
    question = question_service.create_question(question_data)
    
    if question:
        return JsonResponse({
            'success': True,
            'message': 'Question created successfully',
            'question_id': question.id
        })
    else:
        return JsonResponse({'error': 'Failed to create question'}, status=500)


@teacher_required
@require_http_methods(["POST"])
def question_edit_view(request, question_id):
    """
    Edit an existing question.
    """
    question = get_object_or_404(Question, pk=question_id)
    exam = question.exam
    
    # Check if user owns this exam
    teacher = auth_service.get_current_teacher(request)
    if exam.created_by.pk != teacher.pk:
        return JsonResponse({'error': 'Permission denied'}, status=403)
    
    # Get form data
    question_type = request.POST.get('question_type')
    question_text = request.POST.get('question_text', '').strip()
    points = request.POST.get('points', '1.0')
    
    # Validation
    if not question_text:
        return JsonResponse({'error': 'Question text is required'}, status=400)
    
    try:
        points = float(points)
        if points <= 0:
            return JsonResponse({'error': 'Points must be positive'}, status=400)
    except ValueError:
        return JsonResponse({'error': 'Points must be a valid number'}, status=400)
    
    # Prepare question data based on type
    question_data = {
        'question_type': question_type,
        'question_text': question_text,
        'points': points
    }
    
    # Handle type-specific data
    if question_type == QuestionType.MCQ:
        # Parse MCQ options
        options = []
        correct_answer = request.POST.get('correct_answer', '')
        option_keys = request.POST.getlist('option_keys[]')
        option_values = request.POST.getlist('option_values[]')
        
        for key, value in zip(option_keys, option_values):
            if key and value:
                options.append({'key': key, 'value': value})
        
        if len(options) < 2:
            return JsonResponse({'error': 'MCQ must have at least 2 options'}, status=400)
        if not correct_answer:
            return JsonResponse({'error': 'Correct answer is required'}, status=400)
        
        question_data['options'] = options
        question_data['correct_answer'] = correct_answer
    
    elif question_type == QuestionType.IDENTIFICATION:
        correct_answer = request.POST.get('correct_answer', '').strip()
        if not correct_answer:
            return JsonResponse({'error': 'Correct answer is required'}, status=400)
        # Store as list to support multiple acceptable answers
        question_data['correct_answer'] = [correct_answer]
    
    elif question_type == QuestionType.ENUMERATION:
        correct_answers = request.POST.get('correct_answers', '').strip()
        if not correct_answers:
            return JsonResponse({'error': 'Correct answers are required'}, status=400)
        # Split by comma or newline
        answers = [ans.strip() for ans in correct_answers.replace('\n', ',').split(',') if ans.strip()]
        min_required = request.POST.get('min_required', len(answers))
        try:
            min_required = int(min_required)
        except ValueError:
            min_required = len(answers)
        
        question_data['correct_answer'] = {
            'items': answers,
            'min_required': min_required
        }
    
    elif question_type == QuestionType.TRUE_FALSE:
        correct_answer = request.POST.get('correct_answer', '')
        if correct_answer not in ['true', 'false']:
            return JsonResponse({'error': 'Correct answer must be true or false'}, status=400)
        question_data['correct_answer'] = correct_answer == 'true'
    
    elif question_type == QuestionType.ESSAY:
        # Essay questions don't have a correct answer
        question_data['correct_answer'] = None
    
    # Update question
    updated_question = question_service.update_question(question_id, question_data)
    
    if updated_question:
        return JsonResponse({
            'success': True,
            'message': 'Question updated successfully',
            'question_id': updated_question.id
        })
    else:
        return JsonResponse({'error': 'Failed to update question'}, status=500)


@teacher_required
@require_http_methods(["POST"])
def question_delete_view(request, question_id):
    """
    Delete a question.
    """
    question = get_object_or_404(Question, pk=question_id)
    
    # Check if user owns the exam
    teacher = auth_service.get_current_teacher(request)
    if question.exam.created_by.pk != teacher.pk:
        return JsonResponse({'error': 'Permission denied'}, status=403)
    
    exam_id = question.exam.id
    success = question_service.delete_question(question_id)
    
    if success:
        messages.success(request, 'Question deleted successfully')
        return redirect('exam_edit', exam_id=exam_id)
    else:
        messages.error(request, 'Failed to delete question')
        return redirect('exam_edit', exam_id=exam_id)


@teacher_required
@require_http_methods(["POST"])
def ai_generate_questions_view(request, exam_id):
    """Generate questions using AI and add them to an exam."""
    import json as json_module
    from services.ai_generation_service import generate_exam_questions

    exam = get_object_or_404(Exam, pk=exam_id)
    teacher = auth_service.get_current_teacher(request)
    if exam.created_by.pk != teacher.pk:
        return JsonResponse({'error': 'Permission denied'}, status=403)

    try:
        body = json_module.loads(request.body)
    except json_module.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    topic = body.get('topic', '').strip()
    subject = body.get('subject', exam.subject or '').strip()
    type_counts = body.get('type_counts', {})
    difficulty = body.get('difficulty', 'medium')
    grade_level = body.get('grade_level', 'grade_11_12')

    if not topic:
        return JsonResponse({'error': 'Topic is required'}, status=400)
    if not type_counts or not any(v > 0 for v in type_counts.values()):
        return JsonResponse({'error': 'At least one question type count must be greater than 0'}, status=400)

    try:
        questions = generate_exam_questions(topic, subject, type_counts=type_counts, difficulty=difficulty, grade_level=grade_level)
    except ValueError as e:
        return JsonResponse({'error': str(e)}, status=400)

    current_max_order = exam.questions.count()
    created_questions = []

    for i, q in enumerate(questions):
        question = Question.objects.create(
            exam=exam,
            question_type=q['question_type'],
            question_text=q['question_text'],
            options=q.get('options', []),
            correct_answer=q.get('correct_answer') or '',
            points=q.get('points', 1.0),
            order_index=current_max_order + i + 1,
        )
        created_questions.append({
            'id': question.id,
            'question_type': question.question_type,
            'question_text': question.question_text,
            'points': float(question.points),
        })

    return JsonResponse({
        'status': 'ok',
        'message': f'Generated {len(created_questions)} questions',
        'questions': created_questions,
    })


@teacher_required
@require_http_methods(["POST"])
def ai_inline_generate_view(request, exam_id):
    """
    Generate a single question via AI without saving to DB.
    Used for inline question generation in the Add Question modal.
    Returns the question text and metadata for the teacher to review before saving.
    """
    import json as json_module
    from services.ai_generation_service import generate_exam_questions

    exam = get_object_or_404(Exam, pk=exam_id)
    teacher = auth_service.get_current_teacher(request)
    if exam.created_by.pk != teacher.pk:
        return JsonResponse({'error': 'Permission denied'}, status=403)

    try:
        body = json_module.loads(request.body)
    except json_module.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)

    question_type = body.get('question_type', 'MCQ')
    grade_level = body.get('grade_level', 'grade_11_12')
    subject = exam.subject or 'General'

    if question_type not in ['MCQ', 'TRUE_FALSE', 'IDENTIFICATION', 'ENUMERATION', 'ESSAY']:
        return JsonResponse({'error': 'Invalid question type'}, status=400)

    # Build topic context from exam title, description, and existing questions
    topic_parts = [exam.title]
    if exam.description:
        topic_parts.append(exam.description)
    existing_questions = exam.questions.all().order_by('order_index')[:10]
    if existing_questions:
        sample_texts = [q.question_text[:80] for q in existing_questions[:5]]
        topic_parts.append('Related questions already in this exam: ' + '; '.join(sample_texts))
    topic = '. '.join(topic_parts)

    try:
        questions = generate_exam_questions(
            topic=topic,
            subject=subject,
            type_counts={question_type: 1},
            difficulty='medium',
            grade_level=grade_level,
        )
    except ValueError as e:
        return JsonResponse({'error': str(e)}, status=400)

    if not questions:
        return JsonResponse({'error': 'AI failed to generate a question. Please try again.'}, status=500)

    q = questions[0]
    return JsonResponse({
        'success': True,
        'question': {
            'question_type': q['question_type'],
            'question_text': q['question_text'],
            'options': q.get('options', []),
            'correct_answer': q.get('correct_answer'),
            'points': q.get('points', 1.0),
        }
    })


@teacher_required
def exam_takers_view(request, exam_id):
    """
    Display all students who have taken a specific exam.
    Shows student name, ID, scores, and answers for grading.
    """
    from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
    
    exam = get_object_or_404(Exam, pk=exam_id)
    
    # Check if user owns this exam
    teacher = auth_service.get_current_teacher(request)
    if exam.created_by.pk != teacher.pk:
        messages.error(request, 'You do not have permission to view this exam')
        return redirect('exam_list')
    
    # Import necessary models
    from attempts.models import Attempt, AttemptStatus
    from users.models import Student
    
    # Get all attempts for this exam (submitted and graded)
    attempts = Attempt.objects.filter(
        exam_id=exam_id,
        status__in=[AttemptStatus.SUBMITTED, AttemptStatus.GRADED]
    ).select_related('student').order_by('-submitted_at')
    
    # Calculate total possible points
    total_possible_points = sum(float(q.points) for q in exam.questions.all())
    
    # Prepare takers data
    takers_data = []
    for attempt in attempts:
        student = attempt.student
        
        # Calculate percentage
        percentage = 0
        if total_possible_points > 0:
            percentage = (float(attempt.total_score) / total_possible_points) * 100
        
        takers_data.append({
            'attempt': attempt,
            'student': student,
            'percentage': round(percentage, 2),
            'total_possible': total_possible_points
        })
    
    # Implement pagination (20 items per page)
    paginator = Paginator(takers_data, 20)
    page_number = request.GET.get('page', 1)
    
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)
    
    # Build breadcrumbs
    breadcrumbs = build_breadcrumbs(
        ('Dashboard', reverse('teacher_dashboard')),
        ('My Exams', reverse('exam_list')),
        f'{exam.title} - Takers'
    )
    
    context = {
        'exam': exam,
        'takers_data': page_obj.object_list,
        'page_obj': page_obj,
        'total_takers': len(takers_data),
        'total_possible_points': total_possible_points,
        'question_count': exam.questions.count(),
        'page_breadcrumbs': breadcrumbs
    }
    
    return render(request, 'exams/exam_takers.html', context)


@teacher_required
def get_exam_students_view(request, exam_id):
    """
    API endpoint to get list of students from classes assigned to an exam.
    Used by the reopen modal to populate student selection.
    """
    exam = get_object_or_404(Exam, pk=exam_id)
    
    # Check if user owns this exam
    teacher = auth_service.get_current_teacher(request)
    if exam.created_by.pk != teacher.pk:
        return JsonResponse({'error': 'Permission denied'}, status=403)
    
    # Get classes assigned to this exam
    from exams.models import ExamClassAssignment
    assigned_classes = ExamClassAssignment.objects.filter(
        exam=exam
    ).select_related('class_assigned').values_list('class_assigned_id', flat=True)
    
    # Get all students from assigned classes
    from users.models import Student
    students = Student.objects.filter(
        class_assigned_id__in=assigned_classes
    ).order_by('last_name', 'first_name')
    
    # Serialize student data
    students_data = []
    for student in students:
        students_data.append({
            'id': student.id,
            'school_id': student.student_id,
            'full_name': student.get_full_name(),
            'class_name': str(student.class_assigned) if student.class_assigned else 'No Class'
        })
    
    return JsonResponse({
        'success': True,
        'students': students_data,
        'count': len(students_data)
    })


@teacher_required
@ensure_csrf_cookie
def item_summary_view(request, exam_id):
    """
    Display DepEd-style Item Summary Sheet with item analysis.
    Shows difficulty levels, action needed, and competency summary.
    Optionally generates AI-powered teacher analysis.
    """
    from services.item_analysis_service import ItemAnalysisService

    exam = get_object_or_404(Exam, pk=exam_id)

    teacher = auth_service.get_current_teacher(request)
    if exam.created_by.pk != teacher.pk:
        messages.error(request, 'You do not have permission to view this exam')
        return redirect('exam_list')

    service = ItemAnalysisService()
    summary_data = service.get_item_summary(exam_id)

    if summary_data is None:
        messages.error(request, 'Exam not found')
        return redirect('exam_list')

    breadcrumbs = build_breadcrumbs(
        ('Dashboard', reverse('teacher_dashboard')),
        ('My Exams', reverse('exam_list')),
        (exam.title, reverse('exam_takers', args=[exam_id])),
        'Item Summary'
    )

    items = summary_data.get('items') or []
    difficulty_dist = summary_data.get('difficulty_distribution') or {}
    def _dist_count(level):
        entry = difficulty_dist.get(level)
        if isinstance(entry, dict):
            return int(entry.get('count', 0) or 0)
        return int(entry or 0)
    chart_data = {
        'difficulty_labels': ['Easy', 'Moderately Easy', 'Average', 'Difficult', 'Very Difficult'],
        'difficulty_values': [
            _dist_count('Easy'),
            _dist_count('Moderately Easy'),
            _dist_count('Average'),
            _dist_count('Difficult'),
            _dist_count('Very Difficult'),
        ],
        'item_labels': [f"Item {it.get('item_no')}" for it in items],
        'item_percents': [int(it.get('percent_correct') or 0) for it in items],
        'item_types': [it.get('question_type', '') for it in items],
    }

    cached_result = getattr(exam, 'ai_analysis_result', None)
    cached_analysis = cached_result.analysis if cached_result else None
    cached_generated_at = cached_result.generated_at.isoformat() if cached_result and cached_result.generated_at else None
    cached_model_used = cached_result.model_used if cached_result else ''

    item_no_to_question_id = {str(int(it.get('item_no'))): int(it.get('question_id')) for it in items if it.get('item_no') is not None and it.get('question_id') is not None}

    overall_stats = summary_data.get('overall_stats') or {}
    mps_data = summary_data.get('mps_data') or {}
    mastery_level = 'Mastered' if overall_stats.get('avg_percent', 0) >= 75 else ('Nearing Mastery' if overall_stats.get('avg_percent', 0) >= 50 else 'Low Mastery')
    ai_stats = {
        'total_items': summary_data.get('total_items', 0),
        'total_learners': summary_data.get('total_learners', 0),
        'overall_mps': mps_data.get('overall_mps', 0),
        'passing_rate': overall_stats.get('passing_rate', 0),
        'mastery_level': mastery_level,
    }

    context = {
        'exam': exam,
        'summary': summary_data,
        'chart_data': chart_data,
        'ai_stats': ai_stats,
        'item_no_to_question_id_json': item_no_to_question_id,
        'cached_analysis': cached_analysis,
        'cached_generated_at': cached_generated_at,
        'cached_model_used': cached_model_used,
        'page_breadcrumbs': breadcrumbs,
    }

    return render(request, 'exams/item_summary.html', context)


@teacher_required
@require_http_methods(["POST"])
def item_summary_ai_analyze_view(request, exam_id):
    """
    AJAX endpoint to generate AI-powered teacher analysis for item summary.
    Returns JSON with the analysis results. Persists the result so it is available
    on subsequent page loads and survives page refreshes.
    """
    from services.item_analysis_service import ItemAnalysisService
    from exams.models import ItemAnalysisAIResult
    from services.ai_generation_service import get_ai_config

    exam = get_object_or_404(Exam, pk=exam_id)

    teacher = auth_service.get_current_teacher(request)
    if exam.created_by.pk != teacher.pk:
        return JsonResponse({'error': 'Permission denied'}, status=403)

    try:
        service = ItemAnalysisService()
        summary_data = service.get_item_summary(exam_id)

        if not summary_data or not summary_data.get('has_data'):
            return JsonResponse({'error': 'No graded attempts available for analysis'}, status=400)

        analysis = service.generate_ai_analysis(summary_data)
        if analysis is None:
            return JsonResponse({
                'error': 'AI analysis unavailable. Check your AI API settings in Superadmin > AI Settings.'
            }, status=503)

        model_used = (get_ai_config() or {}).get('model', '') or ''
        ItemAnalysisAIResult.objects.update_or_create(
            exam=exam,
            defaults={
                'analysis': analysis,
                'model_used': model_used,
                'generated_by': request.user,
            },
        )

        return JsonResponse({'success': True, 'analysis': analysis, 'model_used': model_used})
    except Exception as e:
        logger.exception('AI analyze failed for exam %s', exam_id)
        return JsonResponse({
            'error': 'AI analysis failed: ' + str(e),
        }, status=500)


@teacher_required
@require_http_methods(["POST"])
def item_summary_ai_clear_view(request, exam_id):
    """AJAX endpoint to delete the cached AI analysis so the next visit is fresh."""
    from exams.models import ItemAnalysisAIResult

    exam = get_object_or_404(Exam, pk=exam_id)

    teacher = auth_service.get_current_teacher(request)
    if exam.created_by.pk != teacher.pk:
        return JsonResponse({'error': 'Permission denied'}, status=403)

    ItemAnalysisAIResult.objects.filter(exam=exam).delete()
    return JsonResponse({'success': True})


@teacher_required
def item_summary_print_view(request, exam_id):
    """
    Render a clean, DepEd-style printable HTML view of the item summary report.
    Returns a self-contained document with @page A4 sizing, page numbers, and
    server-side rendered content (no chart.js dependency) so it prints well on
    any browser or PDF converter.
    """
    from services.item_analysis_service import ItemAnalysisService

    exam = get_object_or_404(Exam, pk=exam_id)

    teacher = auth_service.get_current_teacher(request)
    if exam.created_by.pk != teacher.pk:
        messages.error(request, 'You do not have permission to view this exam')
        return redirect('exam_list')

    service = ItemAnalysisService()
    summary_data = service.get_item_summary(exam_id)

    if summary_data is None:
        messages.error(request, 'Exam not found')
        return redirect('exam_list')

    cached_result = getattr(exam, 'ai_analysis_result', None)

    context = {
        'exam': exam,
        'summary': summary_data,
        'cached_analysis': cached_result.analysis if cached_result else None,
        'cached_generated_at': cached_result.generated_at if cached_result else None,
        'cached_model_used': cached_result.model_used if cached_result else '',
    }

    return render(request, 'exams/item_summary_print.html', context)


@teacher_required
@require_http_methods(["GET"])
def item_summary_export_excel_view(request, exam_id):
    """
    Export the per-exam Item Summary report to Excel with DepEd-style formatting.
    Includes overall statistics, per-item analysis, competency summary, and the
    student-by-item response matrix as the per-exam MPS matrix.
    """
    import os
    import re
    from io import BytesIO
    from django.conf import settings
    from django.http import HttpResponse
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.drawing.image import Image as XlImage
    from services.item_analysis_service import ItemAnalysisService

    exam = get_object_or_404(Exam, pk=exam_id)

    teacher = auth_service.get_current_teacher(request)
    if exam.created_by.pk != teacher.pk:
        messages.error(request, 'You do not have permission to export this exam')
        return redirect('exam_list')

    try:
        service = ItemAnalysisService()
        summary_data = service.get_item_summary(exam_id)

        if not summary_data or not summary_data.get('has_data'):
            messages.warning(request, 'No graded attempts to export for this exam.')
            return redirect('item_summary', exam_id=exam_id)

        overall_stats = summary_data.get('overall_stats') or {}
        items = summary_data.get('items') or []
        mps_data = summary_data.get('mps_data') or {}
        student_matrix = summary_data.get('student_matrix') or {}
        competency_summary = summary_data.get('competency_summary') or []

        cached_result = getattr(exam, 'ai_analysis_result', None)
        cached_analysis = cached_result.analysis if cached_result else None
        cached_generated_at = cached_result.generated_at if cached_result else None
        cached_model_used = cached_result.model_used if cached_result else ''

        wb = Workbook()
        wb.remove(wb.active)

        brand_path = os.path.join(settings.BASE_DIR, 'static', 'img', 'brand.png')
        has_brand = os.path.isfile(brand_path)

        def add_brand_image(worksheet, anchor):
            if not has_brand:
                return False
            try:
                image = XlImage(brand_path)
            except Exception as exc:
                logger.warning('Skipping Item Summary Excel brand image: %s', exc)
                return False
            try:
                image.width = 120
                image.height = 60
                worksheet.add_image(image, anchor)
                return True
            except Exception as exc:
                logger.warning('Failed to attach Item Summary Excel brand image: %s', exc)
                return False

        def safe_export_filename(title):
            cleaned = re.sub(r'[\s/\\:]+', '_', title.strip())
            cleaned = re.sub(r'[^A-Za-z0-9_.-]+', '_', cleaned)
            return cleaned.strip('._-')[:30] or 'Exam'

        def excel_value(value):
            if isinstance(value, str):
                return re.sub(r'[\x00-\x1F]', '', value)
            return value

        def write_cell(worksheet, row_index, column_index, value=None):
            return worksheet.cell(row=row_index, column=column_index, value=excel_value(value))

        def auto_fit_columns(worksheet, start_col, end_col, max_width=42, min_width=10, padding=2):
            """Auto-size columns based on the widest content in each column.
            Excel column width units roughly equal the number of characters that
            fit using the default font (Calibri 11). We add a small padding so
            wrapped headers are not flush against the right border.
            """
            for col_idx in range(start_col, end_col + 1):
                letter = get_column_letter(col_idx)
                longest = 0
                for cell in worksheet[letter]:
                    if cell.value is None:
                        continue
                    val = str(cell.value)
                    if val.startswith('='):
                        continue
                    longest = max(longest, len(val))
                if longest == 0:
                    worksheet.column_dimensions[letter].width = min_width
                    continue
                worksheet.column_dimensions[letter].width = max(min_width, min(max_width, longest + padding))

        title_font = Font(name='Calibri', bold=True, size=14)
        header_font = Font(name='Calibri', bold=True, size=12)
        subheader_font = Font(name='Calibri', bold=True, size=11)
        header_fill = PatternFill(start_color='1F2937', end_color='1F2937', fill_type='solid')
        header_text = Font(name='Calibri', bold=True, size=11, color='FFFFFF')
        green_font = Font(name='Calibri', bold=True, color='15803D')
        yellow_font = Font(name='Calibri', bold=True, color='B45309')
        red_font = Font(name='Calibri', bold=True, color='B91C1C')
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        def get_mps_font(mps_val):
            if mps_val >= 75:
                return green_font
            elif mps_val >= 50:
                return yellow_font
            return red_font

        # --- Sheet: Item Summary ---
        ws = wb.create_sheet(title='Item Summary')
        row = 1

        if add_brand_image(ws, 'A1'):
            row = 5

        ws.cell(row=row, column=1, value='ITEM SUMMARY SHEET')
        ws.cell(row=row, column=1).font = title_font
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)
        row += 1
        ws.cell(row=row, column=1, value='Mean Percentage Score (MPS) — Per Exam')
        ws.cell(row=row, column=1).font = subheader_font
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)
        row += 2

        info_labels = [
            ('Exam:', exam.title),
            ('Subject:', exam.subject or 'Not specified'),
            ('Quarter:', exam.quarter.name if exam.quarter else 'Not specified'),
            ('Teacher:', exam.created_by.user.get_full_name() or exam.created_by.user.username),
            ('No. of Learners:', str(summary_data.get('total_learners', 0))),
            ('Total Items:', str(len(items))),
            ('Total Points:', str(summary_data.get('total_possible', 0))),
        ]
        for label, value in info_labels:
            ws.cell(row=row, column=1, value=label)
            ws.cell(row=row, column=1).font = Font(bold=True)
            write_cell(ws, row, 2, value)
            row += 1
        row += 1

        if overall_stats:
            ws.cell(row=row, column=1, value='OVERALL STATISTICS')
            ws.cell(row=row, column=1).font = header_font
            row += 1
            stat_headers = ['Metric', 'Value']
            for col_idx, h in enumerate(stat_headers, 1):
                cell = ws.cell(row=row, column=col_idx, value=h)
                cell.font = header_text
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center')
                cell.border = thin_border
            row += 1
            stat_rows = [
                ('Average Score', f"{overall_stats.get('avg_score', 0)} / {overall_stats.get('total_possible', 0)} ({overall_stats.get('avg_percent', 0)}%)"),
                ('Passing Rate (60%)', f"{overall_stats.get('passing_rate', 0)}% ({overall_stats.get('passing_count', 0)} passed, {overall_stats.get('failing_count', 0)} failed)"),
                ('Highest Score', str(overall_stats.get('highest_score', 0))),
                ('Lowest Score', str(overall_stats.get('lowest_score', 0))),
                ('Standard Deviation', str(overall_stats.get('std_dev', 0))),
            ]
            for label, value in stat_rows:
                ws.cell(row=row, column=1, value=label).border = thin_border
                write_cell(ws, row, 2, value).border = thin_border
                row += 1
            row += 1

        if mps_data:
            ws.cell(row=row, column=1, value='MEAN PERCENTAGE SCORE (MPS)')
            ws.cell(row=row, column=1).font = header_font
            row += 1
            mps_headers = ['Metric', 'Value']
            for col_idx, h in enumerate(mps_headers, 1):
                cell = ws.cell(row=row, column=col_idx, value=h)
                cell.font = header_text
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center')
                cell.border = thin_border
            row += 1
            mps_rows = [
                ('Overall MPS', f"{mps_data.get('overall_mps', 0)}%"),
                ('Total Correct', str(mps_data.get('total_correct', 0))),
                ('Total Possible Answers', str(mps_data.get('total_possible_answers', 0))),
                ('Number of Learners', str(mps_data.get('total_learners', 0))),
                ('Number of Items', str(mps_data.get('total_items', 0))),
            ]
            for label, value in mps_rows:
                ws.cell(row=row, column=1, value=label).border = thin_border
                cell = write_cell(ws, row, 2, value)
                cell.border = thin_border
                if label == 'Overall MPS':
                    cell.font = get_mps_font(mps_data.get('overall_mps', 0))
                row += 1

            per_class = mps_data.get('per_class') or []
            if per_class:
                row += 1
                ws.cell(row=row, column=1, value='MPS BY CLASS')
                ws.cell(row=row, column=1).font = header_font
                row += 1
                class_headers = ['Class', 'Grade', 'Strand', 'Section', 'Learners', 'Total Correct', 'Total Possible', 'MPS']
                for col_idx, h in enumerate(class_headers, 1):
                    cell = ws.cell(row=row, column=col_idx, value=h)
                    cell.font = header_text
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center')
                    cell.border = thin_border
                row += 1
                for cls_data in per_class:
                    row_data = [
                        cls_data.get('class_name', ''),
                        cls_data.get('grade_level', ''),
                        cls_data.get('strand', ''),
                        cls_data.get('section', ''),
                        cls_data.get('learners', 0),
                        cls_data.get('total_correct', 0),
                        cls_data.get('total_possible', 0),
                        f"{cls_data.get('mps', 0)}%",
                    ]
                    for col_idx, value in enumerate(row_data, 1):
                        cell = write_cell(ws, row, col_idx, value)
                        cell.border = thin_border
                        if col_idx >= 5:
                            cell.alignment = Alignment(horizontal='center')
                        if col_idx == 8:
                            cell.font = get_mps_font(cls_data.get('mps', 0))
                    row += 1
            row += 1

        # --- Item Analysis table ---
        ws.cell(row=row, column=1, value='ITEM ANALYSIS')
        ws.cell(row=row, column=1).font = header_font
        row += 1
        item_headers = ['Item', 'Type', 'Correct', 'Wrong', 'Skipped', '% Correct', 'Difficulty', 'Action']
        for col_idx, h in enumerate(item_headers, 1):
            cell = ws.cell(row=row, column=col_idx, value=h)
            cell.font = header_text
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
        row += 1
        for item in items:
            row_data = [
                item.get('item_no', ''),
                item.get('question_type', ''),
                item.get('num_correct', 0),
                item.get('num_wrong', 0),
                item.get('num_skipped', 0),
                f"{item.get('percent_correct', 0)}%",
                item.get('difficulty_level', ''),
                item.get('action_needed', ''),
            ]
            for col_idx, value in enumerate(row_data, 1):
                cell = write_cell(ws, row, col_idx, value)
                cell.border = thin_border
                if col_idx in (1, 3, 4, 5, 6, 7):
                    cell.alignment = Alignment(horizontal='center')
            row += 1
        row += 1

        # --- Competency Summary (non-tabular, DepEd style) ---
        competency_summary = summary_data.get('competency_summary') or []
        if competency_summary:
            ws.cell(row=row, column=1, value='COMPETENCY SUMMARY')
            ws.cell(row=row, column=1).font = header_font
            row += 1
            comp_label_font = Font(name='Calibri', bold=True, size=11)
            comp_value_font = Font(name='Calibri', size=11)
            for comp in competency_summary:
                competency_text = str(comp.get('competency', ''))
                items_text = str(comp.get('items', ''))
                avg_percent = comp.get('avg_percent', 0)
                mastery_text = str(comp.get('mastery_level', ''))
                intervention_text = str(comp.get('intervention', ''))

                line_cell = ws.cell(row=row, column=1, value=competency_text)
                line_cell.font = comp_label_font
                line_cell.alignment = Alignment(wrap_text=True, vertical='top')
                ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)
                row += 1

                detail_cell = ws.cell(
                    row=row,
                    column=1,
                    value=(
                        '   Items: ' + items_text + '   |   '
                        'Average: ' + f"{avg_percent}%" + '   |   '
                        'Mastery: ' + mastery_text + '   |   '
                        'Intervention: ' + intervention_text
                    ),
                )
                detail_cell.font = comp_value_font
                detail_cell.alignment = Alignment(wrap_text=True, vertical='top', indent=1)
                ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)
                if avg_percent >= 75:
                    detail_cell.font = Font(name='Calibri', size=11, color='15803D', bold=True)
                elif avg_percent >= 50:
                    detail_cell.font = Font(name='Calibri', size=11, color='B45309', bold=True)
                else:
                    detail_cell.font = Font(name='Calibri', size=11, color='B91C1C', bold=True)
                row += 1
                row += 1

        auto_fit_columns(ws, 1, 8, max_width=42, min_width=10, padding=2)

        row += 1
        ws.cell(row=row, column=1, value='Prepared by:').font = Font(bold=True, size=10)
        row += 1
        ws.cell(row=row, column=1, value='________________________________________').font = Font(size=10)
        row += 1
        ws.cell(row=row, column=1, value='Name of Teacher:').font = Font(bold=True, size=10)
        write_cell(ws, row, 2, exam.created_by.user.get_full_name() or exam.created_by.user.username)
        ws.cell(row=row, column=2).font = Font(bold=True, size=10)

        # --- Sheet: Student-by-Item Matrix ---
        if student_matrix and student_matrix.get('students'):
            ws2 = wb.create_sheet(title='Student-Item Matrix')
            row = 1

            if add_brand_image(ws2, 'A1'):
                row = 5

            ws2.cell(row=row, column=1, value='STUDENT-BY-ITEM RESPONSE MATRIX')
            ws2.cell(row=row, column=1).font = title_font
            ws2.merge_cells(start_row=row, start_column=1, end_row=row, end_column=student_matrix['total_items'] + 5)
            row += 1
            write_cell(ws2, row, 1, f'Exam: {exam.title}')
            ws2.cell(row=row, column=1).font = subheader_font
            row += 1
            write_cell(ws2, row, 1, f'No. of Learners: {student_matrix["total_learners"]} | No. of Items: {student_matrix["total_items"]} | MPS: {mps_data.get("overall_mps", 0)}%')
            row += 2

            total_items_count = student_matrix['total_items']
            ws2.cell(row=row, column=1, value='#')
            ws2.cell(row=row, column=2, value='Student Name')
            ws2.cell(row=row, column=3, value='Class')
            for cell in (
                ws2.cell(row=row, column=1),
                ws2.cell(row=row, column=2),
                ws2.cell(row=row, column=3),
            ):
                cell.font = header_text
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center')
                cell.border = thin_border
            for item_idx, item in enumerate(items):
                cell = ws2.cell(row=row, column=item_idx + 4, value=item.get('item_no', ''))
                cell.font = header_text
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center')
                cell.border = thin_border
            total_col = total_items_count + 4
            pct_col = total_items_count + 5
            for col_idx, label in [(total_col, 'Total'), (pct_col, '%')]:
                cell = ws2.cell(row=row, column=col_idx, value=label)
                cell.font = header_text
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center')
                cell.border = thin_border
            row += 1

            green_fill = PatternFill(start_color='DCFCE7', end_color='DCFCE7', fill_type='solid')
            red_fill = PatternFill(start_color='FEE2E2', end_color='FEE2E2', fill_type='solid')

            for idx, student in enumerate(student_matrix['students'], 1):
                ws2.cell(row=row, column=1, value=idx)
                ws2.cell(row=row, column=1).border = thin_border
                ws2.cell(row=row, column=1).alignment = Alignment(horizontal='center')

                write_cell(ws2, row, 2, student.get('name', ''))
                ws2.cell(row=row, column=2).border = thin_border

                write_cell(ws2, row, 3, student.get('class_name', ''))
                ws2.cell(row=row, column=3).border = thin_border

                for item_idx, mark in enumerate(student.get('responses', [])):
                    col = item_idx + 4
                    cell = ws2.cell(row=row, column=col, value=mark)
                    cell.alignment = Alignment(horizontal='center')
                    cell.border = thin_border
                    cell.fill = green_fill if mark == 1 else red_fill

                cell_total = ws2.cell(row=row, column=total_col, value=student.get('total_correct', 0))
                cell_total.font = Font(bold=True)
                cell_total.alignment = Alignment(horizontal='center')
                cell_total.border = thin_border

                cell_pct = ws2.cell(row=row, column=pct_col, value=f"{student.get('percent', 0)}%")
                cell_pct.alignment = Alignment(horizontal='center')
                cell_pct.border = thin_border
                if student.get('percent', 0) >= 75:
                    cell_pct.font = green_font
                elif student.get('percent', 0) < 50:
                    cell_pct.font = red_font

                row += 1

            footer_fill = PatternFill(start_color='E5E7EB', end_color='E5E7EB', fill_type='solid')
            ws2.cell(row=row, column=2, value='Total Correct').font = Font(bold=True)
            ws2.cell(row=row, column=2).border = thin_border
            ws2.cell(row=row, column=2).fill = footer_fill

            for item_idx, count in enumerate(student_matrix.get('per_item_correct', [])):
                col = item_idx + 4
                cell = ws2.cell(row=row, column=col, value=count)
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal='center')
                cell.border = thin_border
                cell.fill = footer_fill

            cell = ws2.cell(row=row, column=total_col, value=mps_data.get('total_correct', 0))
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
            cell.fill = footer_fill
            row += 1

            ws2.cell(row=row, column=2, value='% Correct').font = Font(bold=True)
            ws2.cell(row=row, column=2).border = thin_border
            ws2.cell(row=row, column=2).fill = footer_fill

            for item_idx, pct in enumerate(student_matrix.get('per_item_percent', [])):
                col = item_idx + 4
                cell = ws2.cell(row=row, column=col, value=f"{pct}%")
                cell.alignment = Alignment(horizontal='center')
                cell.border = thin_border
                cell.fill = footer_fill
                if pct >= 75:
                    cell.font = green_font
                elif pct < 50:
                    cell.font = red_font

            cell = ws2.cell(row=row, column=pct_col, value=f"{mps_data.get('overall_mps', 0)}%")
            cell.font = Font(bold=True, color='1F2937')
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
            cell.fill = footer_fill
            row += 2

            ws2.cell(
                row=row,
                column=1,
                value=f'MPS = ({mps_data.get("total_correct", 0)} / {mps_data.get("total_possible_answers", 0)}) x 100 = {mps_data.get("overall_mps", 0)}%'
            ).font = Font(bold=True, size=12)

            ws2.column_dimensions['A'].width = 5
            ws2.column_dimensions['B'].width = 28
            ws2.column_dimensions['C'].width = 16
            for i in range(4, total_items_count + 6):
                ws2.column_dimensions[get_column_letter(i)].width = 5
            ws2.column_dimensions[get_column_letter(total_col)].width = 8
            ws2.column_dimensions[get_column_letter(pct_col)].width = 8

        # --- Sheet: AI Analysis ---
        if cached_analysis:
            ws3 = wb.create_sheet(title='AI Analysis')
            row = 1
            if add_brand_image(ws3, 'A1'):
                row = 5

            ws3.cell(row=row, column=1, value='AI TEACHER\u2019S ANALYSIS')
            ws3.cell(row=row, column=1).font = title_font
            ws3.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
            row += 1
            ws3.cell(row=row, column=1, value='DepEd-aligned insights, intervention plan, and revision recommendations')
            ws3.cell(row=row, column=1).font = subheader_font
            ws3.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
            row += 1
            if cached_generated_at or cached_model_used:
                meta_parts = []
                if cached_generated_at:
                    meta_parts.append('Generated: ' + cached_generated_at.strftime('%B %d, %Y %I:%M %p'))
                if cached_model_used:
                    meta_parts.append('Model: ' + cached_model_used)
                ws3.cell(row=row, column=1, value=' \u2022 '.join(meta_parts))
                ws3.cell(row=row, column=1).font = Font(name='Calibri', size=10, color='6B7280', italic=True)
                ws3.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
                row += 1
            row += 1

            ai_purple_fill = PatternFill(start_color='6D28D9', end_color='6D28D9', fill_type='solid')
            ai_purple_text = Font(name='Calibri', bold=True, size=11, color='FFFFFF')
            ai_section_fill = PatternFill(start_color='EDE9FE', end_color='EDE9FE', fill_type='solid')
            ai_overall_fill = PatternFill(start_color='FAF5FF', end_color='FAF5FF', fill_type='solid')
            ai_green_fill = PatternFill(start_color='F0FDF4', end_color='F0FDF4', fill_type='solid')
            ai_orange_fill = PatternFill(start_color='FFF7ED', end_color='FFF7ED', fill_type='solid')

            if cached_analysis.get('overall_assessment'):
                ws3.cell(row=row, column=1, value='Overall Assessment')
                ws3.cell(row=row, column=1).font = Font(name='Calibri', bold=True, size=11, color='4C1D95')
                ws3.cell(row=row, column=1).fill = ai_section_fill
                ws3.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
                row += 1
                cell = ws3.cell(row=row, column=1, value=cached_analysis['overall_assessment'])
                cell.font = Font(name='Calibri', size=11, italic=True)
                cell.fill = ai_overall_fill
                cell.alignment = Alignment(wrap_text=True, vertical='top')
                ws3.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
                ws3.row_dimensions[row].height = 60
                row += 2

            strengths = cached_analysis.get('strengths') or []
            areas = cached_analysis.get('areas_for_improvement') or []
            if strengths or areas:
                col_strength = 1
                col_areas = 3
                if strengths:
                    cell = ws3.cell(row=row, column=col_strength, value='Strengths (' + str(len(strengths)) + ')')
                    cell.font = Font(name='Calibri', bold=True, size=11, color='166534')
                    cell.fill = ai_green_fill
                    ws3.merge_cells(start_row=row, start_column=col_strength, end_row=row, end_column=col_strength + 1)
                if areas:
                    cell = ws3.cell(row=row, column=col_areas, value='Areas for Improvement (' + str(len(areas)) + ')')
                    cell.font = Font(name='Calibri', bold=True, size=11, color='9A3412')
                    cell.fill = ai_orange_fill
                    ws3.merge_cells(start_row=row, start_column=col_areas, end_row=row, end_column=col_areas + 1)
                row += 1
                max_len = max(len(strengths), len(areas), 1)
                for i in range(max_len):
                    if i < len(strengths):
                        cell = ws3.cell(row=row, column=col_strength, value='\u2022 ' + strengths[i])
                        cell.fill = ai_green_fill
                        cell.alignment = Alignment(wrap_text=True, vertical='top')
                        ws3.merge_cells(start_row=row, start_column=col_strength, end_row=row, end_column=col_strength + 1)
                    if i < len(areas):
                        cell = ws3.cell(row=row, column=col_areas, value='\u2022 ' + areas[i])
                        cell.fill = ai_orange_fill
                        cell.alignment = Alignment(wrap_text=True, vertical='top')
                        ws3.merge_cells(start_row=row, start_column=col_areas, end_row=row, end_column=col_areas + 1)
                    ws3.row_dimensions[row].height = 30
                    row += 1
                row += 1

            steps = cached_analysis.get('intervention_plan') or []
            if steps:
                ws3.cell(row=row, column=1, value='Intervention Plan (' + str(len(steps)) + ' step' + ('s' if len(steps) != 1 else '') + ')')
                ws3.cell(row=row, column=1).font = Font(name='Calibri', bold=True, size=11, color='1F2937')
                ws3.cell(row=row, column=1).fill = ai_section_fill
                ws3.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
                row += 1
                for i, step in enumerate(steps, 1):
                    cell = ws3.cell(row=row, column=1, value=str(i) + '.')
                    cell.font = Font(name='Calibri', bold=True, size=11, color='FFFFFF')
                    cell.fill = ai_purple_fill
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    cell2 = ws3.cell(row=row, column=2, value=step)
                    cell2.font = Font(name='Calibri', size=11)
                    cell2.alignment = Alignment(wrap_text=True, vertical='center')
                    ws3.merge_cells(start_row=row, start_column=2, end_row=row, end_column=4)
                    ws3.row_dimensions[row].height = 28
                    row += 1
                row += 1

            items_to_revise = cached_analysis.get('items_to_revise') or []
            if items_to_revise:
                ws3.cell(row=row, column=1, value='Items Recommended for Revision (' + str(len(items_to_revise)) + ')')
                ws3.cell(row=row, column=1).font = Font(name='Calibri', bold=True, size=11, color='991B1B')
                ws3.cell(row=row, column=1).fill = ai_section_fill
                ws3.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
                row += 1
                chips_text = '   '.join('\u2022 ' + str(it) for it in items_to_revise)
                cell = ws3.cell(row=row, column=1, value=chips_text)
                cell.font = Font(name='Calibri', size=11, bold=True, color='991B1B')
                cell.alignment = Alignment(horizontal='left', vertical='center')
                ws3.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
                ws3.row_dimensions[row].height = 24
                row += 1
                row += 1

            cell = ws3.cell(row=row, column=1, value='Note: AI suggestions are guidance only. Always validate against your own classroom context and DepEd curriculum guides before applying.')
            cell.font = Font(name='Calibri', size=9, italic=True, color='6B7280')
            ws3.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
            row += 2

            teacher_name = exam.created_by.user.get_full_name() or exam.created_by.user.username
            sig_top = row
            cell_l = ws3.cell(row=sig_top, column=1, value=teacher_name)
            cell_l.font = Font(name='Calibri', bold=True, size=11)
            cell_l.alignment = Alignment(horizontal='center')
            ws3.merge_cells(start_row=sig_top, start_column=1, end_row=sig_top, end_column=2)
            cell_r = ws3.cell(row=sig_top, column=3, value='______________________________')
            cell_r.font = Font(name='Calibri', bold=True, size=11)
            cell_r.alignment = Alignment(horizontal='center')
            ws3.merge_cells(start_row=sig_top, start_column=3, end_row=sig_top, end_column=4)
            row += 1
            cell_l = ws3.cell(row=row, column=1, value='Subject Teacher')
            cell_l.font = Font(name='Calibri', size=9, color='6B7280')
            cell_l.alignment = Alignment(horizontal='center')
            ws3.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
            cell_r = ws3.cell(row=row, column=3, value='School Head / Department Head')
            cell_r.font = Font(name='Calibri', size=9, color='6B7280')
            cell_r.alignment = Alignment(horizontal='center')
            ws3.merge_cells(start_row=row, start_column=3, end_row=row, end_column=4)

            ws3.column_dimensions['A'].width = 24
            ws3.column_dimensions['B'].width = 24
            ws3.column_dimensions['C'].width = 24
            ws3.column_dimensions['D'].width = 24

        # Write to response
        output = BytesIO()
        wb.save(output)
        output.seek(0)

        safe_title = safe_export_filename(exam.title)
        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="Item_Summary_{safe_title}.xlsx"'
        return response
    except Exception:
        logger.exception('Failed to build Item Summary Excel export for exam %s', exam_id)
        messages.error(request, 'Unable to generate Excel export. Please try again or contact support.')
        return redirect('item_summary', exam_id=exam_id)


@teacher_required
@require_http_methods(["GET"])
def item_summary_export_word_view(request, exam_id):
    """
    Export the per-exam Item Summary report to Word (DOCX) with DepEd-style formatting.
    Includes overall statistics, per-item analysis, competency summary, and the
    student-by-item response matrix as the per-exam MPS matrix.
    """
    import os
    import re
    from io import BytesIO
    from django.conf import settings
    from django.http import HttpResponse
    from docx import Document
    from docx.shared import Cm, Mm, Pt, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from services.item_analysis_service import ItemAnalysisService

    exam = get_object_or_404(Exam, pk=exam_id)

    teacher = auth_service.get_current_teacher(request)
    if exam.created_by.pk != teacher.pk:
        messages.error(request, 'You do not have permission to export this exam')
        return redirect('exam_list')

    try:
        service = ItemAnalysisService()
        summary_data = service.get_item_summary(exam_id)

        if not summary_data or not summary_data.get('has_data'):
            messages.warning(request, 'No graded attempts to export for this exam.')
            return redirect('item_summary', exam_id=exam_id)

        overall_stats = summary_data.get('overall_stats') or {}
        items = summary_data.get('items') or []
        mps_data = summary_data.get('mps_data') or {}
        competency_summary = summary_data.get('competency_summary') or []
        student_matrix = summary_data.get('student_matrix') or {}

        cached_result = getattr(exam, 'ai_analysis_result', None)
        cached_analysis = cached_result.analysis if cached_result else None
        cached_generated_at = cached_result.generated_at if cached_result else None
        cached_model_used = cached_result.model_used if cached_result else ''

        doc = Document()

        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        _configure_docx_branding(
            doc,
            'ITEM SUMMARY SHEET',
            exam,
            [
                ('Subject', exam.subject or 'Not specified'),
                ('Quarter', exam.quarter.name if exam.quarter else 'Not specified'),
                ('Teacher', exam.created_by.user.get_full_name() or exam.created_by.user.username),
                ('No. of Learners', summary_data.get('total_learners', 0)),
                ('Total Items', len(items)),
            ],
            'Item Summary',
        )

        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(297)
        section.page_height = Mm(210)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        from docx.oxml.ns import qn

        def set_cell_shading(cell, color):
            shading_elm = cell._element.get_or_add_tcPr()
            shading = shading_elm.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): color,
            })
            shading_elm.append(shading)

        def get_mps_color(mps_val):
            if mps_val >= 75:
                return RGBColor(0x15, 0x80, 0x3D)
            elif mps_val >= 50:
                return RGBColor(0xB4, 0x53, 0x09)
            return RGBColor(0xB9, 0x1C, 0x1C)

        # --- DepEd-style header ---
        header_lines = [
            ('Republic of the Philippines', 10, False, RGBColor(0x37, 0x41, 0x51)),
            ('Department of Education', 10, True, RGBColor(0x37, 0x41, 0x51)),
            ('ValuateAI Exam System', 9, False, RGBColor(0x6B, 0x72, 0x80)),
        ]
        for text, size, bold, color in header_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(text)
            run.font.size = Pt(size)
            run.bold = bold
            run.font.color.rgb = color

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(8)
        title.paragraph_format.space_after = Pt(0)
        title_run = title.add_run('ITEM SUMMARY SHEET')
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_before = Pt(0)
        subtitle.paragraph_format.space_after = Pt(2)
        subtitle_run = subtitle.add_run('Mean Percentage Score (MPS) — Per Exam')
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(11)
        subtitle_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        meta_info = doc.add_paragraph()
        meta_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_info.paragraph_format.space_after = Pt(8)
        meta_info_run = meta_info.add_run(
            f"{exam.title}   |   {exam.subject or 'Not specified'}   |   "
            f"{exam.quarter.name if exam.quarter else 'N/A'}   |   "
            f"SY {exam.school_year if hasattr(exam, 'school_year') and exam.school_year else 'N/A'}"
        )
        meta_info_run.font.size = Pt(9)
        meta_info_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        divider = doc.add_paragraph()
        divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        divider.paragraph_format.space_after = Pt(8)
        divider_run = divider.add_run('_' * 72)
        divider_run.font.size = Pt(8)
        divider_run.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)

        doc.add_paragraph()

        if overall_stats:
            doc.add_heading('I. Overall Statistics', level=2)
            stats_table = doc.add_table(rows=1, cols=2)
            stats_table.style = 'Table Grid'
            stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            header_cells = stats_table.rows[0].cells
            header_cells[0].text = 'Metric'
            header_cells[1].text = 'Value'
            for cell in header_cells:
                set_cell_shading(cell, '1F2937')
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(10)

            stats_rows = [
                ('Average Score', f"{overall_stats.get('avg_score', 0)} / {overall_stats.get('total_possible', 0)} ({overall_stats.get('avg_percent', 0)}%)"),
                ('Passing Rate (60%)', f"{overall_stats.get('passing_rate', 0)}% ({overall_stats.get('passing_count', 0)} passed, {overall_stats.get('failing_count', 0)} failed)"),
                ('Highest Score', str(overall_stats.get('highest_score', 0))),
                ('Lowest Score', str(overall_stats.get('lowest_score', 0))),
                ('Standard Deviation', str(overall_stats.get('std_dev', 0))),
            ]
            for label, value in stats_rows:
                row = stats_table.add_row()
                row.cells[0].text = label
                row.cells[1].text = value
                for run in row.cells[0].paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(10)
                for run in row.cells[1].paragraphs[0].runs:
                    run.font.size = Pt(10)
            doc.add_paragraph()

        if mps_data:
            doc.add_heading('II. Mean Percentage Score (MPS)', level=2)
            mps_table = doc.add_table(rows=1, cols=2)
            mps_table.style = 'Table Grid'
            mps_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            mps_table.rows[0].cells[0].text = 'Metric'
            mps_table.rows[0].cells[1].text = 'Value'
            for cell in mps_table.rows[0].cells:
                set_cell_shading(cell, '1F2937')
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(10)

            mps_rows = [
                ('Overall MPS', f"{mps_data.get('overall_mps', 0)}%"),
                ('Total Correct', str(mps_data.get('total_correct', 0))),
                ('Total Possible Answers', str(mps_data.get('total_possible_answers', 0))),
                ('Number of Learners', str(mps_data.get('total_learners', 0))),
                ('Number of Items', str(mps_data.get('total_items', 0))),
            ]
            for label, value in mps_rows:
                row = mps_table.add_row()
                row.cells[0].text = label
                row.cells[1].text = value
                for run in row.cells[0].paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(10)
                for run in row.cells[1].paragraphs[0].runs:
                    run.font.size = Pt(10)
                    if label == 'Overall MPS':
                        run.bold = True
                        run.font.color.rgb = get_mps_color(mps_data.get('overall_mps', 0))

            per_class = mps_data.get('per_class') or []
            if per_class:
                doc.add_paragraph()
                doc.add_heading('MPS by Class', level=3)
                class_table = doc.add_table(rows=1, cols=8)
                class_table.style = 'Table Grid'
                class_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                class_headers = ['Class', 'Grade', 'Strand', 'Section', 'Learners', 'Total Correct', 'Total Possible', 'MPS']
                for idx, h in enumerate(class_headers):
                    cell = class_table.rows[0].cells[idx]
                    cell.text = h
                    set_cell_shading(cell, '1F2937')
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.size = Pt(9)
                for cls_data in per_class:
                    row = class_table.add_row()
                    row_data = [
                        str(cls_data.get('class_name', '')),
                        str(cls_data.get('grade_level', '')),
                        str(cls_data.get('strand', '')),
                        str(cls_data.get('section', '')),
                        str(cls_data.get('learners', 0)),
                        str(cls_data.get('total_correct', 0)),
                        str(cls_data.get('total_possible', 0)),
                        f"{cls_data.get('mps', 0)}%",
                    ]
                    for idx, value in enumerate(row_data):
                        cell = row.cells[idx]
                        cell.text = value
                        for para in cell.paragraphs:
                            if idx > 0:
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in para.runs:
                                run.font.size = Pt(9)
                                if idx == 7:
                                    run.bold = True
                                    run.font.color.rgb = get_mps_color(cls_data.get('mps', 0))
            doc.add_paragraph()

        # Item Analysis
        doc.add_heading('III. Item Analysis', level=2)
        item_table = doc.add_table(rows=1, cols=8)
        item_table.style = 'Table Grid'
        item_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        item_headers = ['Item', 'Type', 'Correct', 'Wrong', 'Skipped', '% Correct', 'Difficulty', 'Action']
        for idx, h in enumerate(item_headers):
            cell = item_table.rows[0].cells[idx]
            cell.text = h
            set_cell_shading(cell, '1F2937')
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(9)
        for item in items:
            row = item_table.add_row()
            row_data = [
                str(item.get('item_no', '')),
                str(item.get('question_type', '')),
                str(item.get('num_correct', 0)),
                str(item.get('num_wrong', 0)),
                str(item.get('num_skipped', 0)),
                f"{item.get('percent_correct', 0)}%",
                str(item.get('difficulty_level', '')),
                str(item.get('action_needed', '')),
            ]
            for idx, value in enumerate(row_data):
                cell = row.cells[idx]
                cell.text = value
                for para in cell.paragraphs:
                    if idx in (0, 2, 3, 4, 5, 6):
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(9)
        doc.add_paragraph()

        # Competency Summary (non-tabular, DepEd style)
        if competency_summary:
            doc.add_heading('IV. Competency Summary', level=2)
            for comp in competency_summary:
                competency_text = str(comp.get('competency', ''))
                items_text = str(comp.get('items', ''))
                avg_percent = comp.get('avg_percent', 0)
                mastery_text = str(comp.get('mastery_level', ''))
                intervention_text = str(comp.get('intervention', ''))

                heading_para = doc.add_paragraph()
                heading_para.paragraph_format.space_after = Pt(2)
                heading_para.paragraph_format.space_before = Pt(6)
                heading_run = heading_para.add_run(competency_text)
                heading_run.bold = True
                heading_run.font.size = Pt(11)

                detail_para = doc.add_paragraph()
                detail_para.paragraph_format.space_after = Pt(2)
                detail_para.paragraph_format.left_indent = Pt(18)
                detail_run = detail_para.add_run(
                    'Items: ' + items_text + '   |   '
                    'Average: ' + f"{avg_percent}%" + '   |   '
                    'Mastery: ' + mastery_text + '   |   '
                    'Intervention: ' + intervention_text
                )
                detail_run.font.size = Pt(10)
                detail_run.bold = True
                detail_run.font.color.rgb = get_mps_color(avg_percent)
            doc.add_paragraph()

        # Student-by-Item Matrix
        if student_matrix and student_matrix.get('students'):
            doc.add_page_break()
            doc.add_heading('V. Student-by-Item Response Matrix', level=2)

            matrix_info = doc.add_paragraph()
            matrix_info.add_run('Legend: ').bold = True
            matrix_info.add_run('1 = Correct, 0 = Wrong/Skipped')

            total_items_count = student_matrix['total_items']
            max_items_per_table = 25

            for chunk_idx, start_item in enumerate(range(0, total_items_count, max_items_per_table)):
                end_item = min(start_item + max_items_per_table, total_items_count)
                chunk_size = end_item - start_item

                if chunk_idx > 0:
                    cont_para = doc.add_paragraph()
                    cont_para.add_run(f'(Continued - Items {start_item + 1} to {end_item})').italic = True

                cols_in_chunk = chunk_size + 5
                matrix_table = doc.add_table(rows=1, cols=cols_in_chunk)
                matrix_table.style = 'Table Grid'
                matrix_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                matrix_table.cell(0, 0).text = '#'
                matrix_table.cell(0, 1).text = 'Student Name'
                matrix_table.cell(0, 2).text = 'Class'
                for c in (0, 1, 2):
                    set_cell_shading(matrix_table.cell(0, c), '1F2937')
                    for para in matrix_table.cell(0, c).paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.size = Pt(7)

                for i in range(start_item + 1, end_item + 1):
                    cell = matrix_table.cell(0, i - start_item + 2)
                    cell.text = str(i)
                    set_cell_shading(cell, '1F2937')
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.size = Pt(7)

                matrix_table.cell(0, chunk_size + 3).text = 'Total'
                matrix_table.cell(0, chunk_size + 4).text = '%'
                for c in (chunk_size + 3, chunk_size + 4):
                    set_cell_shading(matrix_table.cell(0, c), '1F2937')
                    for para in matrix_table.cell(0, c).paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.size = Pt(7)

                for s_idx, student in enumerate(student_matrix['students'], 1):
                    row = matrix_table.add_row()
                    row.cells[0].text = str(s_idx)
                    row.cells[1].text = student.get('name', '')
                    row.cells[2].text = student.get('class_name', '')

                    for item_idx in range(start_item, end_item):
                        col = item_idx - start_item + 3
                        mark = student['responses'][item_idx]
                        cell = row.cells[col]
                        cell.text = str(mark)
                        if mark == 1:
                            set_cell_shading(cell, 'DCFCE7')
                        else:
                            set_cell_shading(cell, 'FEE2E2')

                    row.cells[chunk_size + 3].text = str(student.get('total_correct', 0))
                    row.cells[chunk_size + 4].text = f"{student.get('percent', 0)}%"

                    for cell in row.cells:
                        for para in cell.paragraphs:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in para.runs:
                                run.font.size = Pt(7)
                    for para in row.cells[1].paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for para in row.cells[2].paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                footer_row = matrix_table.add_row()
                footer_row.cells[0].text = ''
                footer_row.cells[1].text = 'Total'
                footer_row.cells[2].text = ''
                for item_idx in range(start_item, end_item):
                    col = item_idx - start_item + 3
                    footer_row.cells[col].text = str(student_matrix['per_item_correct'][item_idx])
                footer_row.cells[chunk_size + 3].text = str(mps_data.get('total_correct', 0))
                footer_row.cells[chunk_size + 4].text = f"{mps_data.get('overall_mps', 0)}%"

                for cell in footer_row.cells:
                    set_cell_shading(cell, 'E5E7EB')
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.bold = True
                            run.font.size = Pt(7)

        # --- VI. AI Teacher's Analysis ---
        if cached_analysis:
            doc.add_page_break()
            doc.add_heading('VI. AI Teacher\u2019s Analysis', level=2)

            meta_parts = []
            if cached_generated_at:
                meta_parts.append('Generated: ' + cached_generated_at.strftime('%B %d, %Y %I:%M %p'))
            if cached_model_used:
                meta_parts.append('Model: ' + cached_model_used)
            if meta_parts:
                meta_para = doc.add_paragraph()
                meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                meta_run = meta_para.add_run(' \u2022 '.join(meta_parts))
                meta_run.italic = True
                meta_run.font.size = Pt(9)
                meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

            if cached_analysis.get('overall_assessment'):
                overall_table = doc.add_table(rows=1, cols=1)
                overall_table.style = 'Table Grid'
                overall_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = overall_table.rows[0].cells[0]
                set_cell_shading(cell, 'FAF5FF')
                p = cell.paragraphs[0]
                run_label = p.add_run('Overall Assessment\n')
                run_label.bold = True
                run_label.font.size = Pt(11)
                run_label.font.color.rgb = RGBColor(0x4C, 0x1D, 0x95)
                run_text = p.add_run(cached_analysis['overall_assessment'])
                run_text.italic = True
                run_text.font.size = Pt(11)
                doc.add_paragraph()

            strengths = cached_analysis.get('strengths') or []
            areas = cached_analysis.get('areas_for_improvement') or []
            if strengths or areas:
                sa_table = doc.add_table(rows=1, cols=2)
                sa_table.style = 'Table Grid'
                sa_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for idx, (title_text, items, fill_color, title_color) in enumerate([
                    ('Strengths (' + str(len(strengths)) + ')', strengths, 'F0FDF4', RGBColor(0x16, 0x65, 0x34)),
                    ('Areas for Improvement (' + str(len(areas)) + ')', areas, 'FFF7ED', RGBColor(0x9A, 0x34, 0x12)),
                ]):
                    cell = sa_table.rows[0].cells[idx]
                    set_cell_shading(cell, fill_color)
                    cell.text = ''
                    p_title = cell.paragraphs[0]
                    run_t = p_title.add_run(title_text + '\n')
                    run_t.bold = True
                    run_t.font.size = Pt(11)
                    run_t.font.color.rgb = title_color
                    for s in items:
                        p_item = cell.add_paragraph()
                        run_i = p_item.add_run('\u2022 ' + s)
                        run_i.font.size = Pt(10)
                doc.add_paragraph()

            steps = cached_analysis.get('intervention_plan') or []
            if steps:
                doc.add_heading('Intervention Plan', level=3)
                step_table = doc.add_table(rows=len(steps), cols=2)
                step_table.style = 'Table Grid'
                step_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for i, step in enumerate(steps, 1):
                    num_cell = step_table.rows[i - 1].cells[0]
                    num_cell.text = str(i)
                    set_cell_shading(num_cell, '6D28D9')
                    for para in num_cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.bold = True
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    text_cell = step_table.rows[i - 1].cells[1]
                    text_cell.text = step
                    for para in text_cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(10)
                doc.add_paragraph()

            items_to_revise = cached_analysis.get('items_to_revise') or []
            if items_to_revise:
                revise_table = doc.add_table(rows=1, cols=1)
                revise_table.style = 'Table Grid'
                revise_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = revise_table.rows[0].cells[0]
                set_cell_shading(cell, 'FEF2F2')
                p = cell.paragraphs[0]
                run_t = p.add_run('Items Recommended for Revision (' + str(len(items_to_revise)) + ')\n')
                run_t.bold = True
                run_t.font.size = Pt(11)
                run_t.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
                chips = '   '.join('\u2022 ' + str(it) for it in items_to_revise)
                run_c = p.add_run(chips)
                run_c.font.size = Pt(11)
                run_c.bold = True
                run_c.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
                doc.add_paragraph()

            note_para = doc.add_paragraph()
            note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            note_run = note_para.add_run('Note: AI suggestions are guidance only. Always validate against your own classroom context and DepEd curriculum guides before applying.')
            note_run.italic = True
            note_run.font.size = Pt(9)
            note_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        # --- Signatures block (2-column) ---
        doc.add_paragraph()
        teacher_name = exam.created_by.user.get_full_name() or exam.created_by.user.username
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c in range(2):
            sig_table.cell(0, c).text = '\n______________________________________________'
            sig_table.cell(0, c).paragraphs[0].runs[0].font.size = Pt(11)
            sig_table.cell(0, c).paragraphs[0].runs[0].bold = True
        sig_table.cell(0, 0).text = teacher_name + '\n______________________________________________'
        sig_table.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(11)
        sig_table.cell(0, 0).paragraphs[0].runs[0].bold = True
        sig_table.cell(1, 0).text = 'Subject Teacher'
        sig_table.cell(1, 1).text = 'School Head / Department Head'
        for c in range(2):
            for para in sig_table.cell(1, c).paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        safe_title = re.sub(r'[\s/\\:]+', '_', (exam.title or '').strip())
        safe_title = re.sub(r'[^A-Za-z0-9_.-]+', '_', safe_title).strip('._-')[:30] or 'Exam'
        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="Item_Summary_{safe_title}.docx"'
        return response
    except Exception:
        logger.exception('Failed to build Item Summary Word export for exam %s', exam_id)
        messages.error(request, 'Unable to generate Word export. Please try again or contact support.')
        return redirect('item_summary', exam_id=exam_id)


@teacher_required
def mps_quarter_list_view(request):
    """
    List all quarters that have exams for the current teacher, with their
    overall MPS shown as a summary card per quarter.
    """
    from services.item_analysis_service import ItemAnalysisService

    teacher = auth_service.get_current_teacher(request)
    service = ItemAnalysisService()
    quarter_summaries = service.get_mps_quarter_list(teacher)

    breadcrumbs = build_breadcrumbs(
        ('Dashboard', reverse('teacher_dashboard')),
        ('My Exams', reverse('exam_list')),
        'MPS Report',
    )

    context = {
        'quarter_summaries': quarter_summaries,
        'has_data': any(q.get('has_data') for q in quarter_summaries),
        'page_breadcrumbs': breadcrumbs,
    }

    return render(request, 'exams/mps_quarter_list.html', context)


@teacher_required
def mps_quarter_detail_view(request, quarter_id):
    """
    Display a per-quarter DepEd-style MPS (Mean Percentage Score) report.
    Aggregates all exams belonging to the selected quarter for the teacher.
    """
    from services.item_analysis_service import ItemAnalysisService
    from users.models import Quarter

    quarter = get_object_or_404(Quarter, pk=quarter_id)

    teacher = auth_service.get_current_teacher(request)
    service = ItemAnalysisService()
    summary_data = service.get_mps_quarter_summary(quarter, teacher)

    if not summary_data:
        messages.error(request, 'Quarter not found')
        return redirect('mps_quarter_list')

    breadcrumbs = build_breadcrumbs(
        ('Dashboard', reverse('teacher_dashboard')),
        ('My Exams', reverse('exam_list')),
        ('MPS Report', reverse('mps_quarter_list')),
        (quarter.name, None),
    )

    mps_gauge_offset = 377
    quarter_summary = summary_data.get('quarter_summary')
    if summary_data.get('has_data') and quarter_summary:
        mps_gauge_offset = round(377 - (377 * quarter_summary['overall_mps'] / 100))

    mps_chart_data = {
        'labels': [],
        'values': [],
        'colors': [],
        'borders': [],
    }
    if quarter_summary and quarter_summary.get('exams'):
        for ex in quarter_summary['exams']:
            if not ex.get('has_data'):
                continue
            label = ex.get('title') or 'Untitled'
            if ex.get('subject'):
                label = label + ' (' + ex['subject'] + ')'
            mps_chart_data['labels'].append(label)
            mps = int(ex.get('overall_mps') or 0)
            mps_chart_data['values'].append(mps)
            if mps >= 75:
                mps_chart_data['colors'].append('rgba(34, 197, 94, 0.85)')
                mps_chart_data['borders'].append('#16a34a')
            elif mps >= 50:
                mps_chart_data['colors'].append('rgba(234, 179, 8, 0.85)')
                mps_chart_data['borders'].append('#ca8a04')
            else:
                mps_chart_data['colors'].append('rgba(239, 68, 68, 0.85)')
                mps_chart_data['borders'].append('#dc2626')
        mps_chart_data['overall_mps'] = int(quarter_summary.get('overall_mps') or 0)

    context = {
        'quarter': quarter,
        'summary': summary_data,
        'mps_gauge_offset': mps_gauge_offset,
        'mps_chart_data': mps_chart_data,
        'page_breadcrumbs': breadcrumbs,
    }

    return render(request, 'exams/mps_quarter_detail.html', context)


def _configure_docx_branding(document, report_title, exam, meta_lines, footer_label):
    import os
    from django.conf import settings
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Inches, Pt, RGBColor

    brand_path = os.path.join(settings.BASE_DIR, 'static', 'img', 'brand.png')

    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(0.2)
    section.footer_distance = Cm(0.2)

    header = section.header
    footer = section.footer
    header.is_linked_to_previous = False
    footer.is_linked_to_previous = False

    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_paragraph.paragraph_format.space_before = Pt(0)
    header_paragraph.paragraph_format.space_after = Pt(0)
    if os.path.isfile(brand_path):
        logo_run = header_paragraph.add_run()
        logo_run.add_picture(brand_path, width=Inches(1.05))
    else:
        placeholder_run = header_paragraph.add_run('School Logo')
        placeholder_run.bold = True
        placeholder_run.font.size = Pt(9)
        placeholder_run.font.color.rgb = RGBColor(90, 90, 90)
    if os.path.isfile(brand_path):
        placeholder_paragraph = header.add_paragraph()
        placeholder_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        placeholder_paragraph.paragraph_format.space_before = Pt(0)
        placeholder_paragraph.paragraph_format.space_after = Pt(0)
        placeholder_run = placeholder_paragraph.add_run('School Logo')
        placeholder_run.bold = True
        placeholder_run.font.size = Pt(8)
        placeholder_run.font.color.rgb = RGBColor(90, 90, 90)

    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)
    footer_run = footer_paragraph.add_run('Generated by: valuateai.onrender.com')
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(107, 114, 128)


def _write_docx_question_block(document, question, question_number):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(2)
    title_paragraph.add_run(f'{question_number}. ').bold = True
    question_run = title_paragraph.add_run(question.question_text)
    question_run.font.size = Pt(11)
    question_run.font.color.rgb = RGBColor(31, 41, 55)

    meta_paragraph = document.add_paragraph()
    meta_paragraph.paragraph_format.space_after = Pt(4)
    meta_paragraph.paragraph_format.left_indent = Cm(0.3)
    meta_run = meta_paragraph.add_run(
        f'{question.get_question_type_display()} • {question.points} point{"s" if question.points != 1 else ""}'
    )
    meta_run.italic = True
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(107, 114, 128)

    if question.question_type == QuestionType.MCQ:
        for option in question.options or []:
            option_paragraph = document.add_paragraph()
            option_paragraph.paragraph_format.left_indent = Cm(0.9)
            option_paragraph.paragraph_format.space_after = Pt(1)
            option_run = option_paragraph.add_run(f"{option.get('key', '')}. {option.get('value', '')}")
            option_run.font.size = Pt(10)
    elif question.question_type == QuestionType.TRUE_FALSE:
        for label in ('True', 'False'):
            option_paragraph = document.add_paragraph()
            option_paragraph.paragraph_format.left_indent = Cm(0.9)
            option_paragraph.paragraph_format.space_after = Pt(1)
            option_run = option_paragraph.add_run(f'☐ {label}')
            option_run.font.size = Pt(10)
    else:
        answer_paragraph = document.add_paragraph()
        answer_paragraph.paragraph_format.left_indent = Cm(0.9)
        answer_paragraph.paragraph_format.space_after = Pt(4)
        answer_run = answer_paragraph.add_run('Answer: ______________________________')
        answer_run.font.size = Pt(10)


@teacher_required
@require_http_methods(["GET"])
def exam_export_word_view(request, exam_id):
    """
    Export a printable exam paper as a Word document.
    Includes exam metadata, a branded header/footer, and the full question list.
    """
    import os
    from io import BytesIO
    from django.conf import settings
    from django.http import HttpResponse
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    exam = get_object_or_404(Exam, pk=exam_id)

    teacher = auth_service.get_current_teacher(request)
    if exam.created_by.pk != teacher.pk:
        messages.error(request, 'Permission denied')
        return redirect('exam_list')

    questions = list(question_service.get_questions_by_exam(exam_id))

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    _configure_docx_branding(
        doc,
        'EXAMINATION PAPER',
        exam,
        [
            ('Exam Title', exam.title),
            ('Subject', exam.subject or 'Not specified'),
            ('Duration', f'{exam.duration_minutes} minutes'),
            ('Teacher', exam.created_by.user.get_full_name() or exam.created_by.user.username),
        ],
        'Exam Export',
    )

    intro_title = doc.add_paragraph()
    intro_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro_run = intro_title.add_run('EXAMINATION PAPER')
    intro_run.bold = True
    intro_run.font.size = Pt(16)

    exam_meta = doc.add_table(rows=4, cols=2)
    exam_meta.style = 'Table Grid'
    meta_rows = [
        ('Exam Title', exam.title),
        ('Subject', exam.subject or 'Not specified'),
        ('Duration', f'{exam.duration_minutes} minutes'),
        ('Instructions', 'Read each item carefully and answer all questions.'),
    ]
    for row_index, (label, value) in enumerate(meta_rows):
        row = exam_meta.rows[row_index]
        row.cells[0].text = label
        row.cells[1].text = value
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.size = Pt(10)

    doc.add_paragraph()
    questions_title = doc.add_paragraph()
    questions_title.paragraph_format.space_after = Pt(6)
    questions_run = questions_title.add_run('Questions')
    questions_run.bold = True
    questions_run.font.size = Pt(13)

    if questions:
        for index, question in enumerate(questions, start=1):
            _write_docx_question_block(doc, question, index)
    else:
        empty_paragraph = doc.add_paragraph()
        empty_paragraph.add_run('No questions have been added to this exam yet.')

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    safe_title = exam.title.replace(' ', '_')[:30]
    response = HttpResponse(
        output.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="Exam_{safe_title}.docx"'
    return response


@teacher_required
@require_http_methods(["GET"])
def mps_quarter_export_excel_view(request, quarter_id):
    """
    Export the per-quarter MPS report to Excel with DepEd-style formatting.
    Includes overall MPS, per-exam breakdown, and item-level data.
    """
    import os
    import re
    from io import BytesIO
    from django.conf import settings
    from django.http import HttpResponse
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.drawing.image import Image as XlImage
    from services.item_analysis_service import ItemAnalysisService
    from users.models import Quarter

    quarter = get_object_or_404(Quarter, pk=quarter_id)

    teacher = auth_service.get_current_teacher(request)
    service = ItemAnalysisService()
    summary_data = service.get_mps_quarter_summary(quarter, teacher)

    if not summary_data or not summary_data.get('has_data'):
        messages.warning(request, 'No graded attempts to export for this quarter.')
        return redirect('mps_quarter_detail', quarter_id=quarter_id)

    representative_exam = summary_data.get('representative_exam')
    if representative_exam is None:
        messages.warning(request, 'No exams in this quarter yet.')
        return redirect('mps_quarter_list')

    quarter_summary = summary_data.get('quarter_summary')
    quarter_matrix = summary_data.get('quarter_matrix')

    wb = Workbook()
    wb.remove(wb.active)

    brand_path = os.path.join(settings.BASE_DIR, 'static', 'img', 'brand.png')
    has_brand = os.path.isfile(brand_path)

    def add_brand_image(worksheet, anchor):
        if not has_brand:
            return False
        try:
            image = XlImage(brand_path)
        except (OSError, ValueError) as exc:
            logger.warning('Skipping MPS Excel brand image: %s', exc)
            return False
        image.width = 120
        image.height = 60
        worksheet.add_image(image, anchor)
        return True

    def safe_export_filename(title):
        cleaned = re.sub(r'[\s/\\:]+', '_', title.strip())
        cleaned = re.sub(r'[^A-Za-z0-9_.-]+', '_', cleaned)
        return cleaned.strip('._-')[:30] or 'Exam'

    def excel_value(value):
        if isinstance(value, str):
            return re.sub(r'[\x00-\x1F]', '', value)
        return value

    def write_cell(worksheet, row_index, column_index, value=None):
        return worksheet.cell(row=row_index, column=column_index, value=excel_value(value))

    title_font = Font(name='Calibri', bold=True, size=14)
    header_font = Font(name='Calibri', bold=True, size=12)
    subheader_font = Font(name='Calibri', bold=True, size=11)
    header_fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
    header_text = Font(name='Calibri', bold=True, size=11, color='FFFFFF')
    green_font = Font(name='Calibri', bold=True, color='006100')
    yellow_font = Font(name='Calibri', bold=True, color='7F6000')
    red_font = Font(name='Calibri', bold=True, color='9C0006')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    def get_mps_font(mps_val):
        if mps_val >= 75:
            return green_font
        elif mps_val >= 50:
            return yellow_font
        return red_font

    def get_performance_level(mps_val):
        if mps_val >= 75:
            return 'Mastered'
        elif mps_val >= 50:
            return 'Nearing Mastery'
        return 'Low Mastery'

    # --- Sheet 0: Quarter Summary ---
    ws = wb.create_sheet(title='Quarter Summary')
    row = 1

    if add_brand_image(ws, 'A1'):
        row = 5

    ws.cell(row=row, column=1, value='QUARTER MEAN PERCENTAGE SCORE (MPS) REPORT')
    ws.cell(row=row, column=1).font = title_font
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=7)
    row += 2

    if quarter_summary:
        quarter_exam_subjects = sorted({
            exam_summary['subject']
            for exam_summary in quarter_summary.get('exams', [])
            if exam_summary.get('subject')
        })
        subject_label = ', '.join(quarter_exam_subjects) if quarter_exam_subjects else 'Not specified'

        info_labels = [
            ('Quarter:', quarter_summary['quarter_name']),
            ('Subject:', subject_label),
            ('Teacher:', representative_exam.created_by.user.get_full_name() or representative_exam.created_by.user.username),
            ('No. of Learners:', str(quarter_summary['graded_attempts'])),
            ('Quarter Exam Count:', str(quarter_summary['exam_count'])),
            ('Quarter MPS:', f"{quarter_summary['overall_mps']}%"),
        ]
        for label, value in info_labels:
            ws.cell(row=row, column=1, value=label)
            ws.cell(row=row, column=1).font = Font(bold=True)
            write_cell(ws, row, 2, value)
            row += 1

        row += 1
        ws.cell(row=row, column=1, value='QUARTER SUMMARY')
        ws.cell(row=row, column=1).font = header_font
        row += 1

        quarter_headers = ['Exam', 'Subject', 'Attempts', 'Items', 'MPS']
        for col_idx, h in enumerate(quarter_headers, 1):
            cell = ws.cell(row=row, column=col_idx, value=h)
            cell.font = header_text
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
        row += 1

        for exam_summary in quarter_summary['exams']:
            row_data = [
                exam_summary['title'],
                exam_summary['subject'],
                exam_summary['total_learners'],
                exam_summary['total_items'],
                f"{exam_summary['overall_mps']}%" if exam_summary['has_data'] else 'No data',
            ]
            for col_idx, value in enumerate(row_data, 1):
                cell = write_cell(ws, row, col_idx, value)
                cell.border = thin_border
                if col_idx in (3, 4, 5):
                    cell.alignment = Alignment(horizontal='center')
                if col_idx == 5 and exam_summary['has_data']:
                    cell.font = get_mps_font(exam_summary['overall_mps'])
            row += 1

        row += 1
        ws.cell(row=row, column=1, value='MPS BY EXAM')
        ws.cell(row=row, column=1).font = header_font
        row += 1

        exam_headers = ['Exam', 'Learners', 'Total Correct', 'Total Possible', 'MPS', 'Level']
        for col_idx, h in enumerate(exam_headers, 1):
            cell = ws.cell(row=row, column=col_idx, value=h)
            cell.font = header_text
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
        row += 1

        for exam_summary in quarter_summary['exams']:
            row_data = [
                exam_summary['title'],
                exam_summary['total_learners'],
                exam_summary['total_correct'],
                exam_summary['total_possible_answers'],
                f"{exam_summary['overall_mps']}%" if exam_summary['has_data'] else 'No data',
                get_performance_level(exam_summary['overall_mps']),
            ]
            for col_idx, value in enumerate(row_data, 1):
                cell = write_cell(ws, row, col_idx, value)
                cell.border = thin_border
                if col_idx in (2, 3, 4, 5, 6):
                    cell.alignment = Alignment(horizontal='center')
                if col_idx == 5 and exam_summary['has_data']:
                    cell.font = get_mps_font(exam_summary['overall_mps'])
            row += 1

        for i, width in enumerate([28, 14, 14, 14, 12, 16], 1):
            ws.column_dimensions[get_column_letter(i)].width = width

    row += 2
    ws.cell(row=row, column=1, value='Prepared by:')
    ws.cell(row=row, column=1).font = Font(bold=True, size=10)
    row += 1
    ws.cell(row=row, column=1, value='________________________________________')
    ws.cell(row=row, column=1).font = Font(size=10)
    row += 1
    ws.cell(row=row, column=1, value='Name of Teacher:')
    ws.cell(row=row, column=1).font = Font(bold=True, size=10)
    write_cell(ws, row, 2, representative_exam.created_by.user.get_full_name() or representative_exam.created_by.user.username)
    ws.cell(row=row, column=2).font = Font(bold=True, size=10)

    # --- Sheet 2: Quarter Matrix ---
    if quarter_matrix and quarter_matrix['students']:
        ws3 = wb.create_sheet(title='Quarter Matrix')
        row = 1

        if add_brand_image(ws3, 'A1'):
            row = 5

        ws3.cell(row=row, column=1, value='QUARTER STUDENT-BY-ITEM RESPONSE MATRIX')
        ws3.cell(row=row, column=1).font = title_font
        ws3.merge_cells(start_row=row, start_column=1, end_row=row, end_column=quarter_matrix['total_items'] + 4)
        row += 1
        write_cell(ws3, row, 1, f'Quarter: {quarter_summary["quarter_name"] if quarter_summary else "Not specified"}')
        ws3.cell(row=row, column=1).font = subheader_font
        row += 1
        write_cell(ws3, row, 1, f'No. of Learners: {quarter_matrix["total_learners"]} | No. of Items: {quarter_matrix["total_items"]}')
        row += 2

        total_items = quarter_matrix['total_items']
        header_row_1 = row
        header_row_2 = row + 1

        ws3.cell(row=header_row_1, column=1, value='#')
        ws3.cell(row=header_row_1, column=2, value='Student Name')
        for cell in (ws3.cell(row=header_row_1, column=1), ws3.cell(row=header_row_1, column=2)):
            cell.font = header_text
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
        ws3.merge_cells(start_row=header_row_1, start_column=1, end_row=header_row_2, end_column=1)
        ws3.merge_cells(start_row=header_row_1, start_column=2, end_row=header_row_2, end_column=2)

        col_cursor = 3
        for section in quarter_matrix['exam_sections']:
            if section['item_count'] <= 0:
                continue
            start_col = col_cursor
            end_col = col_cursor + section['item_count'] - 1
            section_title = f"{section['title']} ({section['overall_mps']}% MPS)"
            write_cell(ws3, header_row_1, start_col, section_title)
            ws3.merge_cells(start_row=header_row_1, start_column=start_col, end_row=header_row_1, end_column=end_col)
            top_cell = ws3.cell(row=header_row_1, column=start_col)
            top_cell.font = header_text
            top_cell.fill = header_fill
            top_cell.alignment = Alignment(horizontal='center', wrap_text=True)
            top_cell.border = thin_border
            col_cursor += section['item_count']

        total_col = total_items + 3
        pct_col = total_items + 4
        for col_idx, label in [(total_col, 'Total'), (pct_col, '%')]:
            ws3.cell(row=header_row_1, column=col_idx, value=label)
            ws3.cell(row=header_row_1, column=col_idx).font = header_text
            ws3.cell(row=header_row_1, column=col_idx).fill = header_fill
            ws3.cell(row=header_row_1, column=col_idx).alignment = Alignment(horizontal='center')
            ws3.cell(row=header_row_1, column=col_idx).border = thin_border
            ws3.merge_cells(start_row=header_row_1, start_column=col_idx, end_row=header_row_2, end_column=col_idx)

        for item_idx, item in enumerate(quarter_matrix['items']):
            cell = ws3.cell(row=header_row_2, column=item_idx + 3, value=item['exam_item_no'])
            cell.font = header_text
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border

        row = header_row_2 + 1
        green_fill = PatternFill(start_color='E2EFDA', end_color='E2EFDA', fill_type='solid')
        red_fill = PatternFill(start_color='FCE4EC', end_color='FCE4EC', fill_type='solid')

        for idx, student in enumerate(quarter_matrix['students'], 1):
            ws3.cell(row=row, column=1, value=idx)
            ws3.cell(row=row, column=1).border = thin_border
            ws3.cell(row=row, column=1).alignment = Alignment(horizontal='center')

            write_cell(ws3, row, 2, student['name'])
            ws3.cell(row=row, column=2).border = thin_border

            for item_idx, mark in enumerate(student['responses']):
                col = item_idx + 3
                cell = ws3.cell(row=row, column=col, value=mark)
                cell.alignment = Alignment(horizontal='center')
                cell.border = thin_border
                cell.fill = green_fill if mark == 1 else red_fill

            cell_total = ws3.cell(row=row, column=total_col, value=student['total_correct'])
            cell_total.font = Font(bold=True)
            cell_total.alignment = Alignment(horizontal='center')
            cell_total.border = thin_border

            cell_pct = ws3.cell(row=row, column=pct_col, value=f"{student['percent']}%")
            cell_pct.alignment = Alignment(horizontal='center')
            cell_pct.border = thin_border
            if student['percent'] >= 75:
                cell_pct.font = green_font
            elif student['percent'] < 50:
                cell_pct.font = red_font

            row += 1

        footer_fill = PatternFill(start_color='D6E4F0', end_color='D6E4F0', fill_type='solid')
        ws3.cell(row=row, column=1, value='')
        ws3.cell(row=row, column=2, value='Total Correct')
        ws3.cell(row=row, column=2).font = Font(bold=True)
        ws3.cell(row=row, column=2).border = thin_border
        ws3.cell(row=row, column=2).fill = footer_fill

        for item_idx, count in enumerate(quarter_matrix['per_item_correct']):
            col = item_idx + 3
            cell = ws3.cell(row=row, column=col, value=count)
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
            cell.fill = footer_fill

        cell = ws3.cell(row=row, column=total_col, value=quarter_matrix['total_correct'])
        cell.font = Font(bold=True)
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_border
        cell.fill = footer_fill
        row += 1

        ws3.cell(row=row, column=2, value='% Correct')
        ws3.cell(row=row, column=2).font = Font(bold=True)
        ws3.cell(row=row, column=2).border = thin_border
        ws3.cell(row=row, column=2).fill = footer_fill

        for item_idx, pct in enumerate(quarter_matrix['per_item_percent']):
            col = item_idx + 3
            cell = ws3.cell(row=row, column=col, value=f"{pct}%")
            cell.alignment = Alignment(horizontal='center')
            cell.border = thin_border
            cell.fill = footer_fill
            if pct >= 75:
                cell.font = green_font
            elif pct < 50:
                cell.font = red_font

        cell = ws3.cell(row=row, column=pct_col, value=f"{quarter_matrix['overall_mps']}%")
        cell.font = Font(bold=True, color='1F4E79')
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_border
        cell.fill = footer_fill
        row += 2

        ws3.cell(row=row, column=1, value=f'MPS = ({quarter_matrix["total_correct"]} / {quarter_matrix["total_possible_answers"]}) x 100 = {quarter_matrix["overall_mps"]}%')
        ws3.cell(row=row, column=1).font = Font(bold=True, size=12)

        ws3.column_dimensions['A'].width = 5
        ws3.column_dimensions['B'].width = 28
        for i in range(3, total_items + 5):
            ws3.column_dimensions[get_column_letter(i)].width = 5
        ws3.column_dimensions[get_column_letter(total_col)].width = 8
        ws3.column_dimensions[get_column_letter(pct_col)].width = 8

    # Write to response
    output = BytesIO()
    wb.save(output)
    output.seek(0)

    safe_title = safe_export_filename(quarter.name)
    response = HttpResponse(
        output.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="MPS_Report_{safe_title}.xlsx"'
    return response


@teacher_required
@require_http_methods(["GET"])
def mps_quarter_export_word_view(request, quarter_id):
    """
    Export the per-quarter MPS report to Word (DOCX) with DepEd-style formatting.
    Includes overall MPS, per-exam breakdown, and item-level data.
    """
    from io import BytesIO
    from django.http import HttpResponse
    from docx import Document
    from docx.shared import Cm, Mm, Pt, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from services.item_analysis_service import ItemAnalysisService
    from users.models import Quarter

    quarter = get_object_or_404(Quarter, pk=quarter_id)

    teacher = auth_service.get_current_teacher(request)
    service = ItemAnalysisService()
    summary_data = service.get_mps_quarter_summary(quarter, teacher)

    if not summary_data or not summary_data.get('has_data'):
        messages.warning(request, 'No graded attempts to export for this quarter.')
        return redirect('mps_quarter_detail', quarter_id=quarter_id)

    representative_exam = summary_data.get('representative_exam')
    if representative_exam is None:
        messages.warning(request, 'No exams in this quarter yet.')
        return redirect('mps_quarter_list')

    quarter_summary = summary_data.get('quarter_summary')
    quarter_matrix = summary_data.get('quarter_matrix')

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    _configure_docx_branding(
        doc,
        'QUARTER MEAN PERCENTAGE SCORE (MPS) REPORT',
        representative_exam,
        [
            ('Quarter', quarter_summary['quarter_name'] if quarter_summary else 'Not specified'),
            ('Subject', ', '.join(
                sorted({
                    ex['subject']
                    for ex in quarter_summary.get('exams', [])
                    if ex.get('subject')
                })
            ) if quarter_summary else 'Not specified'),
            ('Teacher', representative_exam.created_by.user.get_full_name() or representative_exam.created_by.user.username),
            ('No. of Learners', quarter_summary['graded_attempts'] if quarter_summary else 0),
            ('Quarter Exams', quarter_summary['exam_count'] if quarter_summary else 0),
        ],
        'Quarter MPS Report',
    )

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    from docx.oxml.ns import qn

    def set_cell_shading(cell, color):
        shading_elm = cell._element.get_or_add_tcPr()
        shading = shading_elm.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): color,
        })
        shading_elm.append(shading)

    def get_performance_level(mps_val):
        if mps_val >= 75:
            return 'Mastered'
        elif mps_val >= 50:
            return 'Nearing Mastery'
        return 'Low Mastery'

    # Title
    title = doc.add_heading('QUARTER MEAN PERCENTAGE SCORE (MPS) REPORT', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Exam info block
    info_block = [
        ('School', '________________________'),
        ('Teacher', representative_exam.created_by.user.get_full_name() or representative_exam.created_by.user.username),
        ('Subject', ', '.join(
            sorted({
                ex['subject']
                for ex in quarter_summary.get('exams', [])
                if ex.get('subject')
            })
        ) if quarter_summary else 'Not specified'),
        ('Quarter', quarter_summary['quarter_name'] if quarter_summary else 'Not specified'),
        ('No. of Learners', str(quarter_summary['graded_attempts'] if quarter_summary else 0)),
        ('No. of Items', str(quarter_summary['total_items'] if quarter_summary else 0)),
    ]
    for label, value in info_block:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f'{label}: ')
        label_run.bold = True
        label_run.font.size = Pt(10)
        value_run = paragraph.add_run(str(value))
        value_run.font.size = Pt(10)

    doc.add_paragraph()

    # Selected quarter summary
    if quarter_summary:
        doc.add_heading(f'Quarter Summary - {quarter_summary["quarter_name"]}', level=2)
        quarter_meta = doc.add_paragraph()
        quarter_meta.paragraph_format.space_after = Pt(2)
        quarter_meta.add_run('Quarter MPS: ').bold = True
        quarter_meta.add_run(f'{quarter_summary["overall_mps"]}%')

        quarter_meta_2 = doc.add_paragraph()
        quarter_meta_2.paragraph_format.space_after = Pt(4)
        quarter_meta_2.add_run('Graded Attempts: ').bold = True
        quarter_meta_2.add_run(str(quarter_summary['graded_attempts']))

        quarter_table = doc.add_table(rows=1, cols=5)
        quarter_table.style = 'Table Grid'
        quarter_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        quarter_headers = ['Exam', 'Subject', 'Attempts', 'Items', 'MPS']
        header_row = quarter_table.rows[0]
        for idx, h in enumerate(quarter_headers):
            cell = header_row.cells[idx]
            cell.text = h
            set_cell_shading(cell, '1F4E79')
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(9)

        for exam_summary in quarter_summary['exams']:
            row = quarter_table.add_row()
            row_data = [
                exam_summary['title'],
                exam_summary['subject'],
                str(exam_summary['total_learners']),
                str(exam_summary['total_items']),
                f"{exam_summary['overall_mps']}%" if exam_summary['has_data'] else 'No data',
            ]
            for idx, value in enumerate(row_data):
                cell = row.cells[idx]
                cell.text = value
                for para in cell.paragraphs:
                    if idx > 0:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(9)
                        if idx == 4 and exam_summary['has_data']:
                            if exam_summary['overall_mps'] >= 75:
                                run.font.color.rgb = RGBColor(0, 97, 0)
                            elif exam_summary['overall_mps'] >= 50:
                                run.font.color.rgb = RGBColor(127, 96, 0)
                            else:
                                run.font.color.rgb = RGBColor(156, 0, 6)

        doc.add_paragraph()
        doc.add_heading('MPS by Exam', level=2)
        exam_table = doc.add_table(rows=1, cols=6)
        exam_table.style = 'Table Grid'
        exam_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['Exam', 'Learners', 'Total Correct', 'Total Possible', 'MPS', 'Performance Level']
        header_row = exam_table.rows[0]
        for idx, h in enumerate(headers):
            cell = header_row.cells[idx]
            cell.text = h
            set_cell_shading(cell, '1F4E79')
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(9)

        for exam_summary in quarter_summary['exams']:
            row = exam_table.add_row()
            row_data = [
                exam_summary['title'],
                str(exam_summary['total_learners']),
                str(exam_summary['total_correct']),
                str(exam_summary['total_possible_answers']),
                f"{exam_summary['overall_mps']}%" if exam_summary['has_data'] else 'No data',
                get_performance_level(exam_summary['overall_mps']),
            ]
            for idx, value in enumerate(row_data):
                cell = row.cells[idx]
                cell.text = value
                for para in cell.paragraphs:
                    if idx > 0:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(9)
                        if idx == 4 and exam_summary['has_data']:
                            run.bold = True
                            if exam_summary['overall_mps'] >= 75:
                                run.font.color.rgb = RGBColor(0, 97, 0)
                            elif exam_summary['overall_mps'] >= 50:
                                run.font.color.rgb = RGBColor(127, 96, 0)
                            else:
                                run.font.color.rgb = RGBColor(156, 0, 6)

    # Quarter matrix
    if quarter_matrix and quarter_matrix['students']:
        doc.add_page_break()
        doc.add_heading('Quarter Student-by-Item Response Matrix', level=2)

        matrix_para = doc.add_paragraph()
        matrix_para.add_run('Legend: ').bold = True
        matrix_para.add_run('1 = Correct, 0 = Wrong/Skipped')

        total_items = quarter_matrix['total_items']
        max_items_per_table = 25
        item_chunks = []
        for start in range(0, total_items, max_items_per_table):
            end = min(start + max_items_per_table, total_items)
            item_chunks.append((start, end))

        for chunk_idx, (start_item, end_item) in enumerate(item_chunks):
            chunk_size = end_item - start_item
            cols_in_chunk = chunk_size + 4

            if chunk_idx > 0:
                doc.add_paragraph()
                cont_para = doc.add_paragraph()
                cont_para.add_run(f'(Continued - Items {start_item + 1} to {end_item})').italic = True

            matrix_table = doc.add_table(rows=2, cols=cols_in_chunk)
            matrix_table.style = 'Table Grid'
            matrix_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Base header cells
            matrix_table.cell(0, 0).text = '#'
            matrix_table.cell(0, 1).text = 'Student Name'
            set_cell_shading(matrix_table.cell(0, 0), '1F4E79')
            set_cell_shading(matrix_table.cell(0, 1), '1F4E79')
            for cell in (matrix_table.cell(0, 0), matrix_table.cell(0, 1)):
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(7)
            matrix_table.cell(1, 0).text = ''
            matrix_table.cell(1, 1).text = ''

            # Item headers
            for i in range(start_item + 1, end_item + 1):
                cell = matrix_table.cell(1, i - start_item + 1)
                cell.text = str(i)
                set_cell_shading(cell, '1F2937')
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(7)

            # Totals columns
            matrix_table.cell(0, chunk_size + 2).text = 'Total'
            matrix_table.cell(0, chunk_size + 3).text = '%'
            for c in (chunk_size + 2, chunk_size + 3):
                set_cell_shading(matrix_table.cell(0, c), '1F4E79')
                for para in matrix_table.cell(0, c).paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(7)
            matrix_table.cell(1, chunk_size + 2).text = ''
            matrix_table.cell(1, chunk_size + 3).text = ''

            # Student rows
            for s_idx, student in enumerate(quarter_matrix['students'], 1):
                row = matrix_table.add_row()
                row.cells[0].text = str(s_idx)
                row.cells[1].text = student['name']

                for item_idx in range(start_item, end_item):
                    col = item_idx - start_item + 2
                    mark = student['responses'][item_idx]
                    cell = row.cells[col]
                    cell.text = str(mark)
                    if mark == 1:
                        set_cell_shading(cell, 'DCFCE7')
                    else:
                        set_cell_shading(cell, 'FCE4EC')

                row.cells[chunk_size + 2].text = str(student['total_correct'])
                row.cells[chunk_size + 3].text = f"{student['percent']}%"

                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.size = Pt(7)
                for para in row.cells[1].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Footer: totals per item
            footer_row = matrix_table.add_row()
            footer_row.cells[0].text = ''
            footer_row.cells[1].text = 'Total'
            for item_idx in range(start_item, end_item):
                col = item_idx - start_item + 2
                footer_row.cells[col].text = str(quarter_matrix['per_item_correct'][item_idx])
            footer_row.cells[chunk_size + 2].text = str(quarter_matrix['total_correct'])
            footer_row.cells[chunk_size + 3].text = f"{quarter_matrix['overall_mps']}%"

            for cell in footer_row.cells:
                set_cell_shading(cell, 'D6E4F0')
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
                        run.font.size = Pt(7)

    doc.add_paragraph()
    prepared_by_para = doc.add_paragraph()
    prepared_by_para.paragraph_format.space_after = Pt(2)
    prepared_by_run = prepared_by_para.add_run('Prepared by:')
    prepared_by_run.bold = True
    prepared_by_run.font.size = Pt(10)

    signature_line = doc.add_paragraph()
    signature_line.paragraph_format.space_after = Pt(2)
    signature_run = signature_line.add_run('________________________________________')
    signature_run.font.size = Pt(10)

    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_after = Pt(0)
    name_label = name_para.add_run('Name of Teacher: ')
    name_label.bold = True
    name_label.font.size = Pt(10)
    name_value = name_para.add_run(representative_exam.created_by.user.get_full_name() or representative_exam.created_by.user.username)
    name_value.bold = True
    name_value.font.size = Pt(10)

    # Write to response
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    safe_title = re.sub(r'[\s/\\:]+', '_', (quarter.name or '').strip())
    safe_title = re.sub(r'[^A-Za-z0-9_.-]+', '_', safe_title).strip('._-')[:30] or 'Exam'
    response = HttpResponse(
        output.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="MPS_Report_{safe_title}.docx"'
    return response
