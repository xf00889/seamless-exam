from io import BytesIO
import html
import json
import re

from django.contrib.auth.models import User
from django.test import TestCase
from django.urls import reverse
from django.utils import timezone
from docx import Document
from openpyxl import load_workbook

from attempts.models import Attempt, AttemptStatus
from exams.models import Exam, Question, QuestionType
from users.models import Class as SchoolClass, Quarter, Student, Teacher
from exams.models import ExamClassAssignment
from services.item_analysis_service import ItemAnalysisService


def _container_text(container):
    parts = []
    for paragraph in container.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _document_text(document):
    parts = []
    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _section_xml(section_part):
    return section_part._element.xml


class ExamEditPageTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username='teacher1',
            password='testpass123',
            first_name='Test',
            last_name='Teacher',
        )
        self.teacher = Teacher.objects.create(user=self.user, department='Science')
        self.exam = Exam.objects.create(
            title='Geology Quiz',
            subject='Science',
            duration_minutes=60,
            created_by=self.teacher,
        )
        self.question = Question.objects.create(
            exam=self.exam,
            question_type=QuestionType.MCQ,
            question_text='What is the main process of rock formation?',
            options=[
                {'key': 'A', 'value': 'Erosion'},
                {'key': 'B', 'value': 'Weathering'},
            ],
            correct_answer='B',
            points=1,
            order_index=0,
        )

        self.client.force_login(self.user)
        session = self.client.session
        session['user_type'] = 'teacher'
        session.save()

    def test_exam_edit_page_embeds_questions_as_a_json_array(self):
        response = self.client.get(reverse('exam_edit', args=[self.exam.id]))

        self.assertEqual(response.status_code, 200)

        match = re.search(
            r'<script id="questions-data" type="application/json">(.*?)</script>',
            response.content.decode('utf-8'),
            re.DOTALL,
        )
        self.assertIsNotNone(match, 'questions-data JSON script was not rendered')

        questions_payload = json.loads(html.unescape(match.group(1)))

        self.assertIsInstance(questions_payload, list)
        self.assertEqual(len(questions_payload), 1)
        self.assertEqual(questions_payload[0]['id'], self.question.id)
        self.assertEqual(questions_payload[0]['text'], self.question.question_text)


class ExamQuarterSelectionTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username='teacher_quarter',
            password='testpass123',
            first_name='Quarter',
            last_name='Teacher',
        )
        self.teacher = Teacher.objects.create(user=self.user, department='Science')
        self.quarter = Quarter.objects.create(name='1st Quarter')
        self.client.force_login(self.user)
        session = self.client.session
        session['user_type'] = 'teacher'
        session.save()

    def test_exam_create_saves_quarter(self):
        response = self.client.post(
            reverse('exam_create'),
            data={
                'generation_method': 'manual',
                'title': 'Quarter Exam',
                'subject': 'Science',
                'quarter': self.quarter.id,
                'description': 'Quarter-based exam',
                'duration_minutes': 60,
                'assigned_classes': [],
            },
        )

        self.assertEqual(response.status_code, 302)
        exam = Exam.objects.get(title='Quarter Exam')
        self.assertEqual(exam.quarter_id, self.quarter.id)

    def test_exam_create_page_lists_existing_quarters(self):
        response = self.client.get(reverse('exam_create'))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '1st Quarter')


class ExamListQuarterGroupingTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username='teacher_list',
            password='testpass123',
            first_name='List',
            last_name='Teacher',
        )
        self.teacher = Teacher.objects.create(user=self.user, department='Science')
        self.quarter_one = Quarter.objects.create(name='1st Quarter', order=1)
        self.quarter_two = Quarter.objects.create(name='2nd Quarter', order=2)
        self.exam_one = Exam.objects.create(
            title='Quarter One Exam',
            subject='Science',
            duration_minutes=45,
            quarter=self.quarter_one,
            created_by=self.teacher,
        )
        self.exam_two = Exam.objects.create(
            title='Quarter Two Exam',
            subject='Science',
            duration_minutes=45,
            quarter=self.quarter_two,
            created_by=self.teacher,
        )
        self.exam_no_quarter = Exam.objects.create(
            title='Ungrouped Exam',
            subject='Science',
            duration_minutes=45,
            created_by=self.teacher,
        )

        self.client.force_login(self.user)
        session = self.client.session
        session['user_type'] = 'teacher'
        session.save()

    def test_exam_list_groups_exams_by_quarter(self):
        response = self.client.get(reverse('exam_list'))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, '1st Quarter')
        self.assertContains(response, '2nd Quarter')
        self.assertContains(response, 'No Quarter')
        self.assertContains(response, self.exam_one.title)
        self.assertContains(response, self.exam_two.title)
        self.assertContains(response, self.exam_no_quarter.title)


class QuarterSummaryServiceTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username='teacher_summary',
            password='testpass123',
            first_name='Quarter',
            last_name='Summary',
        )
        self.teacher = Teacher.objects.create(user=self.user, department='Science')
        self.quarter = Quarter.objects.create(name='1st Quarter', order=1)
        self.other_quarter = Quarter.objects.create(name='2nd Quarter', order=2)
        self.exam_one = Exam.objects.create(
            title='Quarter Exam 1',
            subject='Science',
            duration_minutes=45,
            quarter=self.quarter,
            created_by=self.teacher,
        )
        self.exam_two = Exam.objects.create(
            title='Quarter Exam 2',
            subject='Science',
            duration_minutes=45,
            quarter=self.quarter,
            created_by=self.teacher,
        )
        self.exam_three = Exam.objects.create(
            title='Quarter Exam 3',
            subject='Science',
            duration_minutes=45,
            quarter=self.other_quarter,
            created_by=self.teacher,
        )
        self.student = Student.objects.create(
            school_id='S-3001',
            first_name='Cora',
            last_name='Mendoza',
            password_hash='legacy',
        )

        self.question_one = Question.objects.create(
            exam=self.exam_one,
            question_type=QuestionType.TRUE_FALSE,
            question_text='Water boils at 100 degrees Celsius at sea level.',
            correct_answer=True,
            points=1,
            order_index=0,
        )
        self.question_two = Question.objects.create(
            exam=self.exam_two,
            question_type=QuestionType.TRUE_FALSE,
            question_text='The moon emits its own light.',
            correct_answer=False,
            points=1,
            order_index=0,
        )
        self.question_three = Question.objects.create(
            exam=self.exam_three,
            question_type=QuestionType.TRUE_FALSE,
            question_text='Plants need sunlight to grow.',
            correct_answer=True,
            points=1,
            order_index=0,
        )

        self.attempt_one = Attempt.objects.create(
            student=self.student,
            exam=self.exam_one,
            total_score=1,
            status=AttemptStatus.GRADED,
            submitted_at=timezone.now(),
        )
        self.attempt_two = Attempt.objects.create(
            student=self.student,
            exam=self.exam_two,
            total_score=0,
            status=AttemptStatus.GRADED,
            submitted_at=timezone.now(),
        )
        self.attempt_three = Attempt.objects.create(
            student=self.student,
            exam=self.exam_three,
            total_score=1,
            status=AttemptStatus.GRADED,
            submitted_at=timezone.now(),
        )

        from attempts.models import Answer

        Answer.objects.create(
            attempt=self.attempt_one,
            question=self.question_one,
            answer_text={'value': True},
            is_correct=True,
            points_earned=1,
            graded_at=timezone.now(),
        )
        Answer.objects.create(
            attempt=self.attempt_two,
            question=self.question_two,
            answer_text={'value': True},
            is_correct=False,
            points_earned=0,
            graded_at=timezone.now(),
        )
        Answer.objects.create(
            attempt=self.attempt_three,
            question=self.question_three,
            answer_text={'value': True},
            is_correct=True,
            points_earned=1,
            graded_at=timezone.now(),
        )

        self.client.force_login(self.user)
        session = self.client.session
        session['user_type'] = 'teacher'
        session.save()

    def test_quarter_summary_aggregates_multiple_exams(self):
        summary = ItemAnalysisService().get_item_summary(self.exam_one.id)

        self.assertIn('quarter_summary', summary)
        self.assertTrue(summary['quarter_summary']['has_data'])
        self.assertEqual(summary['quarter_summary']['quarter_name'], '1st Quarter')
        self.assertEqual(summary['quarter_summary']['exam_count'], 2)
        self.assertEqual(summary['quarter_summary']['overall_mps'], 50.0)
        self.assertEqual(len(summary['quarter_summary']['exams']), 2)

    def test_item_summary_page_shows_quarter_summary(self):
        response = self.client.get(reverse('item_summary', args=[self.exam_one.id]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Quarter Summary')
        self.assertContains(response, '1st Quarter')
        self.assertContains(response, self.exam_one.title)
        self.assertContains(response, self.exam_two.title)

    def test_mps_report_page_shows_quarter_summary(self):
        response = self.client.get(reverse('mps_report', args=[self.exam_one.id]))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'Quarter MPS Overview')
        self.assertContains(response, '1st Quarter')
        self.assertContains(response, '2nd Quarter')
        self.assertContains(response, self.exam_one.title)
        self.assertContains(response, self.exam_two.title)
        self.assertContains(response, 'Exam Student-by-Item Response Matrix')
        self.assertContains(response, 'Quarter Student-by-Item Response Matrix')
        self.assertContains(response, 'All Quarters Student-by-Item Response Matrix')
        nav_labels = {link['label']: link for link in response.context['navbar_nav_links']}
        self.assertTrue(nav_labels['MPS']['active'])
        self.assertFalse(nav_labels['Exams']['active'])

    def test_mps_entrypoint_redirects_to_latest_report(self):
        response = self.client.get(reverse('mps_entrypoint'))

        self.assertEqual(response.status_code, 302)
        self.assertEqual(response.url, reverse('mps_report', args=[self.exam_three.id]))



class ExamDetailPageTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username='teacher2',
            password='testpass123',
            first_name='Detail',
            last_name='Teacher',
        )
        self.teacher = Teacher.objects.create(user=self.user, department='Mathematics')
        self.exam = Exam.objects.create(
            title='Algebra Quiz',
            subject='Math',
            duration_minutes=45,
            description='Linear equations and factoring.',
            created_by=self.teacher,
        )
        self.school_class = SchoolClass.objects.create(
            grade_level='Grade 10',
            strand='STEM',
            section='A',
            teacher=self.teacher,
        )
        ExamClassAssignment.objects.create(exam=self.exam, class_assigned=self.school_class)
        self.student = Student.objects.create(
            school_id='S-1001',
            first_name='Ana',
            last_name='Lopez',
            password_hash='legacy',
            class_assigned=self.school_class,
        )
        self.question = Question.objects.create(
            exam=self.exam,
            question_type=QuestionType.TRUE_FALSE,
            question_text='Linear equations always have one solution.',
            correct_answer=True,
            points=2,
            order_index=0,
        )
        Attempt.objects.create(
            student=self.student,
            exam=self.exam,
            total_score=1.50,
            status=AttemptStatus.SUBMITTED,
            submitted_at=timezone.now(),
        )

        self.client.force_login(self.user)
        session = self.client.session
        session['user_type'] = 'teacher'
        session.save()

    def test_exam_detail_page_shows_exam_overview_and_summary(self):
        response = self.client.get(reverse('exam_detail', args=[self.exam.id]))

        self.assertEqual(response.status_code, 200)
        content = response.content.decode('utf-8')

        self.assertContains(response, 'Exam Details: Algebra Quiz')
        self.assertContains(response, 'px-4 py-2 text-sm font-medium')
        self.assertContains(response, 'Edit Exam')
        self.assertContains(response, 'Edit Questions')
        self.assertContains(response, 'View Takers')
        self.assertContains(response, 'Export Word')
        self.assertContains(response, 'Class Summary')
        self.assertContains(response, 'Taker Summary')
        self.assertContains(response, self.question.question_text)
        self.assertContains(response, 'Grade 10 - STEM - A')
        self.assertContains(response, 'Ana Lopez')
        self.assertContains(response, 'Questions')
        self.assertContains(response, 'Accessible Students')
        self.assertContains(response, 'Recent Takers')
        self.assertIn('reopenModal', content)


class ExamExportWordTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username='teacher3',
            password='testpass123',
            first_name='Export',
            last_name='Teacher',
        )
        self.teacher = Teacher.objects.create(user=self.user, department='English')
        self.exam = Exam.objects.create(
            title='Literature Quiz',
            subject='English',
            duration_minutes=30,
            description='Poetry and prose.',
            created_by=self.teacher,
        )
        Question.objects.create(
            exam=self.exam,
            question_type=QuestionType.MCQ,
            question_text='Which literary device compares two unlike things using "like" or "as"?',
            options=[
                {'key': 'A', 'value': 'Metaphor'},
                {'key': 'B', 'value': 'Simile'},
            ],
            correct_answer='B',
            points=1,
            order_index=0,
        )

        self.client.force_login(self.user)
        session = self.client.session
        session['user_type'] = 'teacher'
        session.save()

    def test_exam_export_word_contains_header_footer_and_questions(self):
        response = self.client.get(reverse('exam_export_word', args=[self.exam.id]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )

        document = Document(BytesIO(response.content))
        header_xml = _section_xml(document.sections[0].header)
        footer_xml = _section_xml(document.sections[0].footer)
        body_text = _document_text(document)

        self.assertIn('School Logo', header_xml)
        self.assertNotIn('EXAMINATION PAPER', header_xml)
        self.assertIn('Literature Quiz', body_text)
        self.assertIn('Generated by: seamless.dpdns.org', footer_xml)
        self.assertIn('Literature Quiz', body_text)
        self.assertIn('Which literary device compares two unlike things using "like" or "as"?', body_text)
        self.assertIn('Metaphor', body_text)
        self.assertIn('Simile', body_text)


class MpsExportWordTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username='teacher4',
            password='testpass123',
            first_name='MPS',
            last_name='Teacher',
        )
        self.teacher = Teacher.objects.create(user=self.user, department='Science')
        self.quarter = Quarter.objects.create(name='1st Quarter', order=1)
        self.other_quarter = Quarter.objects.create(name='2nd Quarter', order=2)
        self.exam = Exam.objects.create(
            title='Physics Check',
            subject='Science',
            duration_minutes=45,
            quarter=self.quarter,
            created_by=self.teacher,
        )
        self.other_exam = Exam.objects.create(
            title='Quarter Review',
            subject='Science',
            duration_minutes=30,
            quarter=self.other_quarter,
            created_by=self.teacher,
        )
        self.school_class = SchoolClass.objects.create(
            grade_level='Grade 9',
            strand='STEM',
            section='B',
            teacher=self.teacher,
        )
        ExamClassAssignment.objects.create(exam=self.exam, class_assigned=self.school_class)
        self.student = Student.objects.create(
            school_id='S-2001',
            first_name='Ben',
            last_name='Cruz',
            password_hash='legacy',
            class_assigned=self.school_class,
        )
        self.question = Question.objects.create(
            exam=self.exam,
            question_type=QuestionType.TRUE_FALSE,
            question_text='Energy can be created from nothing.',
            correct_answer=False,
            points=1,
            order_index=0,
        )
        self.other_question = Question.objects.create(
            exam=self.other_exam,
            question_type=QuestionType.TRUE_FALSE,
            question_text='The earth is flat.',
            correct_answer=False,
            points=1,
            order_index=0,
        )
        self.attempt = Attempt.objects.create(
            student=self.student,
            exam=self.exam,
            total_score=0,
            status=AttemptStatus.GRADED,
            submitted_at=timezone.now(),
        )
        self.other_attempt = Attempt.objects.create(
            student=self.student,
            exam=self.other_exam,
            total_score=1,
            status=AttemptStatus.GRADED,
            submitted_at=timezone.now(),
        )
        from attempts.models import Answer

        Answer.objects.create(
            attempt=self.attempt,
            question=self.question,
            answer_text={'value': True},
            is_correct=False,
            points_earned=0,
            graded_at=timezone.now(),
        )
        Answer.objects.create(
            attempt=self.other_attempt,
            question=self.other_question,
            answer_text={'value': False},
            is_correct=True,
            points_earned=1,
            graded_at=timezone.now(),
        )

        self.client.force_login(self.user)
        session = self.client.session
        session['user_type'] = 'teacher'
        session.save()

    def test_mps_export_word_contains_header_and_footer(self):
        response = self.client.get(reverse('mps_export_word', args=[self.exam.id]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response['Content-Type'],
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )

        document = Document(BytesIO(response.content))
        header_xml = _section_xml(document.sections[0].header)
        footer_xml = _section_xml(document.sections[0].footer)

        self.assertIn('School Logo', header_xml)
        self.assertNotIn('QUARTER MEAN PERCENTAGE SCORE (MPS) REPORT', header_xml)
        self.assertIn('Generated by: seamless.dpdns.org', footer_xml)
        self.assertIn('Prepared by:', _document_text(document))
        self.assertIn('Name of Teacher: MPS Teacher', _document_text(document))
        self.assertIn('Quarter MPS Overview', _document_text(document))
        self.assertIn('Quarter Summary - 1st Quarter', _document_text(document))
        self.assertIn('1st Quarter', _document_text(document))
        self.assertIn('2nd Quarter', _document_text(document))
        self.assertIn('MPS by Exam', _document_text(document))
        self.assertIn('Quarter Student-by-Item Response Matrix', _document_text(document))
        self.assertGreaterEqual(len(document.tables), 4)
        self.assertEqual(document.tables[0].cell(0, 0).text, 'Quarter')
        self.assertEqual(document.tables[1].cell(0, 0).text, 'Exam')
        self.assertEqual(document.tables[2].cell(0, 0).text, 'Exam')
        self.assertEqual(document.tables[3].cell(0, 0).text, '#')
        self.assertNotIn('School:', document.tables[1]._element.xml)


class MpsExportExcelTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username='teacher5',
            password='testpass123',
            first_name='Excel',
            last_name='Teacher',
        )
        self.teacher = Teacher.objects.create(user=self.user, department='Science')
        self.quarter = Quarter.objects.create(name='1st Quarter', order=1)
        self.other_quarter = Quarter.objects.create(name='2nd Quarter', order=2)
        self.exam = Exam.objects.create(
            title='Biology Check',
            subject='Science',
            duration_minutes=45,
            quarter=self.quarter,
            created_by=self.teacher,
        )
        self.other_exam = Exam.objects.create(
            title='Biology Review',
            subject='Science',
            duration_minutes=30,
            quarter=self.other_quarter,
            created_by=self.teacher,
        )
        self.school_class = SchoolClass.objects.create(
            grade_level='Grade 8',
            strand='STEM',
            section='C',
            teacher=self.teacher,
        )
        ExamClassAssignment.objects.create(exam=self.exam, class_assigned=self.school_class)
        self.student = Student.objects.create(
            school_id='S-3002',
            first_name='Dina',
            last_name='Reyes',
            password_hash='legacy',
            class_assigned=self.school_class,
        )
        self.question = Question.objects.create(
            exam=self.exam,
            question_type=QuestionType.TRUE_FALSE,
            question_text='Plants make food through photosynthesis.',
            correct_answer=True,
            points=1,
            order_index=0,
        )
        self.other_question = Question.objects.create(
            exam=self.other_exam,
            question_type=QuestionType.TRUE_FALSE,
            question_text='Fish can breathe on land forever.',
            correct_answer=False,
            points=1,
            order_index=0,
        )
        self.attempt = Attempt.objects.create(
            student=self.student,
            exam=self.exam,
            total_score=1,
            status=AttemptStatus.GRADED,
            submitted_at=timezone.now(),
        )
        self.other_attempt = Attempt.objects.create(
            student=self.student,
            exam=self.other_exam,
            total_score=0,
            status=AttemptStatus.GRADED,
            submitted_at=timezone.now(),
        )
        from attempts.models import Answer

        Answer.objects.create(
            attempt=self.attempt,
            question=self.question,
            answer_text={'value': True},
            is_correct=True,
            points_earned=1,
            graded_at=timezone.now(),
        )
        Answer.objects.create(
            attempt=self.other_attempt,
            question=self.other_question,
            answer_text={'value': True},
            is_correct=False,
            points_earned=0,
            graded_at=timezone.now(),
        )

        self.client.force_login(self.user)
        session = self.client.session
        session['user_type'] = 'teacher'
        session.save()

    def test_mps_export_excel_contains_quarter_summary_sheet(self):
        response = self.client.get(reverse('mps_export_excel', args=[self.exam.id]))

        self.assertEqual(response.status_code, 200)
        workbook = load_workbook(BytesIO(response.content))
        self.assertIn('Quarter Overview', workbook.sheetnames)
        self.assertIn('Quarter Summary', workbook.sheetnames)
        self.assertIn('Quarter Matrix', workbook.sheetnames)

        overview_sheet = workbook['Quarter Overview']
        overview_values = {
            cell.value
            for row in overview_sheet.iter_rows()
            for cell in row
            if cell.value is not None
        }
        self.assertIn('QUARTER MPS OVERVIEW', overview_values)
        self.assertIn('1st Quarter', overview_values)
        self.assertIn('2nd Quarter', overview_values)

        summary_sheet = workbook['Quarter Summary']
        summary_values = {
            cell.value
            for row in summary_sheet.iter_rows()
            for cell in row
            if cell.value is not None
        }
        self.assertIn('QUARTER MEAN PERCENTAGE SCORE (MPS) REPORT', summary_values)
        self.assertIn('QUARTER SUMMARY', summary_values)
        self.assertIn('MPS BY EXAM', summary_values)
        self.assertIn('Biology Check', summary_values)
        self.assertNotIn('Biology Review', summary_values)

        matrix_sheet = workbook['Quarter Matrix']
        matrix_values = {
            cell.value
            for row in matrix_sheet.iter_rows()
            for cell in row
            if cell.value is not None
        }
        self.assertIn('QUARTER STUDENT-BY-ITEM RESPONSE MATRIX', matrix_values)
        self.assertIn('Reyes, Dina', matrix_values)
