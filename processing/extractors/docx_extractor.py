"""
DOCX text extraction using python-docx.
"""
from typing import List, Optional
import logging
from .base_extractor import BaseExtractor

logger = logging.getLogger(__name__)


class DOCXExtractor(BaseExtractor):
    """
    Extractor for DOCX documents using python-docx.
    Preserves document structure including paragraphs and lists.
    """
    
    def __init__(self):
        """Initialize DOCX extractor."""
        super().__init__()
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if python-docx is available."""
        try:
            import docx
            self.docx_available = True
        except ImportError:
            self.logger.error("python-docx not available")
            self.docx_available = False
            raise ImportError("python-docx is required for DOCX extraction")
    
    def extract(self, file_path: str) -> str:
        """
        Extract text from DOCX file preserving structure.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text with structure preserved
            
        Raises:
            FileNotFoundError: If file does not exist
            ValueError: If file is not a DOCX
            Exception: For extraction errors
        """
        self.validate_file_exists(file_path)
        self.validate_file_extension(file_path, ['.docx', '.DOCX'])
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_parts.append(table_text)
            
            return '\n'.join(text_parts)
        
        except Exception as e:
            self.logger.error(f"Error extracting DOCX: {e}")
            raise
    
    def _extract_table_text(self, table) -> str:
        """
        Extract text from a table.
        
        Args:
            table: python-docx Table object
            
        Returns:
            Formatted table text
        """
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells.append(cell_text)
            if cells:
                rows.append(' | '.join(cells))
        
        return '\n'.join(rows)
    
    def extract_paragraphs(self, file_path: str) -> List[str]:
        """
        Extract paragraphs as a list.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of paragraph texts
        """
        self.validate_file_exists(file_path)
        self.validate_file_extension(file_path, ['.docx', '.DOCX'])
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            return paragraphs
        
        except Exception as e:
            self.logger.error(f"Error extracting paragraphs: {e}")
            raise
    
    def extract_with_formatting(self, file_path: str) -> List[dict]:
        """
        Extract text with formatting information.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of dictionaries with 'text' and 'style' keys
        """
        self.validate_file_exists(file_path)
        self.validate_file_extension(file_path, ['.docx', '.DOCX'])
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            formatted_content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    formatted_content.append({
                        'text': text,
                        'style': paragraph.style.name if paragraph.style else 'Normal',
                        'is_heading': paragraph.style.name.startswith('Heading') if paragraph.style else False
                    })
            
            return formatted_content
        
        except Exception as e:
            self.logger.error(f"Error extracting with formatting: {e}")
            raise
    
    def extract_headings(self, file_path: str) -> List[dict]:
        """
        Extract only headings from the document.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of dictionaries with 'text' and 'level' keys
        """
        self.validate_file_exists(file_path)
        self.validate_file_extension(file_path, ['.docx', '.DOCX'])
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            headings = []
            
            for paragraph in doc.paragraphs:
                if paragraph.style and paragraph.style.name.startswith('Heading'):
                    text = paragraph.text.strip()
                    if text:
                        # Extract heading level (e.g., "Heading 1" -> 1)
                        level = 1
                        try:
                            level = int(paragraph.style.name.split()[-1])
                        except:
                            pass
                        
                        headings.append({
                            'text': text,
                            'level': level
                        })
            
            return headings
        
        except Exception as e:
            self.logger.error(f"Error extracting headings: {e}")
            raise
    
    def extract_lists(self, file_path: str) -> List[List[str]]:
        """
        Extract lists from the document.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of lists (each inner list is a set of list items)
        """
        self.validate_file_exists(file_path)
        self.validate_file_extension(file_path, ['.docx', '.DOCX'])
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            lists = []
            current_list = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                # Check if paragraph is a list item
                if paragraph.style and 'List' in paragraph.style.name:
                    if text:
                        current_list.append(text)
                else:
                    # End of current list
                    if current_list:
                        lists.append(current_list)
                        current_list = []
            
            # Add last list if exists
            if current_list:
                lists.append(current_list)
            
            return lists
        
        except Exception as e:
            self.logger.error(f"Error extracting lists: {e}")
            raise
